import streamlit as st
import google.generativeai as genai
import json
import re
import docx
from pypdf import PdfReader
import io
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.auth.transport.requests import Request
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
import base64
from email.mime.text import MIMEText
# Y asegúrate también de que estas de antes siguen ahí:
import time
import httplib2
import google_auth_httplib2 
import os      # <-- AÑADE ESTA LÍNEA
import imgkit  # <-- Y AÑADE ESTA LÍNEA TAMBIÉN

# =============================================================================
#           BLOQUE COMPLETO DE CONFIGURACIÓN Y FUNCIONES DE DRIVE
# =============================================================================
st.set_page_config(layout="wide")

SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/gmail.send',
    'https://www.googleapis.com/auth/userinfo.profile',
    'https://www.googleapis.com/auth/userinfo.email',
    'openid',
    'https://www.googleapis.com/auth/documents.readonly' # <-- NUEVO PERMISO
]
CLIENT_CONFIG = {
    "web": {
        "client_id": st.secrets["GOOGLE_CLIENT_ID"],
        "client_secret": st.secrets["GOOGLE_CLIENT_SECRET"],
        "redirect_uris": [st.secrets["GOOGLE_REDIRECT_URI"]],
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
    }
}
ROOT_FOLDER_NAME = "ProyectosLicitaciones"


# =============================================================================
#           BLOQUE DE FUNCIONES DE DRIVE DEFINITIVO (CON LAS 6 FUNCIONES)
# =============================================================================

def find_or_create_folder(service, folder_name, parent_id=None, retries=3):
    """Busca una carpeta. Si no la encuentra, la crea. Incluye reintentos para errores de red."""
    query = f"name = '{folder_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    
    for attempt in range(retries):
        try:
            response = service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
            files = response.get('files', [])
            if files:
                return files[0]['id']
            else:
                file_metadata = {'name': folder_name, 'mimeType': 'application/vnd.google-apps.folder'}
                if parent_id:
                    file_metadata['parents'] = [parent_id]
                folder = service.files().create(body=file_metadata, fields='id').execute()
                st.toast(f"Carpeta '{folder_name}' creada en tu Drive.")
                return folder.get('id')
        except (TimeoutError, httplib2.ServerNotFoundError) as e:
            if attempt < retries - 1:
                st.toast(f"⏳ Error de red con Drive ({type(e).__name__}). Reintentando... ({attempt + 2}/{retries})")
                time.sleep(2 ** attempt)
            else:
                st.error("❌ No se pudo conectar con Google Drive. Por favor, refresca la página.")
                raise
        except Exception as e:
            st.error(f"Ocurrió un error inesperado con Google Drive: {e}")
            raise

def upload_file_to_drive(service, file_object, folder_id, retries=3):
    """Sube un objeto de archivo a una carpeta de Drive, con reintentos."""
    for attempt in range(retries):
        try:
            file_metadata = {'name': file_object.name, 'parents': [folder_id]}
            file_object.seek(0) 
            media = MediaIoBaseUpload(file_object, mimetype=file_object.type, resumable=True)
            file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
            st.toast(f"📄 Archivo '{file_object.name}' guardado en Drive.")
            return file.get('id')
        except (TimeoutError, httplib2.ServerNotFoundError) as e:
            if attempt < retries - 1:
                st.toast(f"⏳ Error de red al subir archivo. Reintentando... ({attempt + 2}/{retries})")
                time.sleep(2 ** attempt)
            else:
                st.error(f"❌ No se pudo subir el archivo '{file_object.name}' tras varios intentos.")
                raise
        except Exception as e:
            st.error(f"Error inesperado al subir archivo: {e}")
            raise

def delete_file_from_drive(service, file_id, retries=3):
    """Elimina un archivo de Drive por su ID, con reintentos."""
    for attempt in range(retries):
        try:
            service.files().delete(fileId=file_id).execute()
            return True
        except (TimeoutError, httplib2.ServerNotFoundError) as e:
            if attempt < retries - 1:
                st.toast(f"⏳ Error de red al eliminar. Reintentando... ({attempt + 2}/{retries})")
                time.sleep(2 ** attempt)
            else:
                st.error(f"❌ No se pudo eliminar el archivo/carpeta tras varios intentos.")
                return False
        except HttpError as error:
            st.error(f"No se pudo eliminar el archivo: {error}")
            return False

def find_file_by_name(service, file_name, folder_id, retries=3):
    """Busca un archivo por nombre dentro de una carpeta, con reintentos."""
    query = f"name = '{file_name}' and '{folder_id}' in parents and trashed = false"
    for attempt in range(retries):
        try:
            response = service.files().list(q=query, spaces='drive', fields='files(id)').execute()
            files = response.get('files', [])
            return files[0]['id'] if files else None
        except (TimeoutError, httplib2.ServerNotFoundError) as e:
            if attempt < retries - 1:
                st.toast(f"⏳ Error de red buscando archivo. Reintentando... ({attempt + 2}/{retries})")
                time.sleep(2 ** attempt)
            else:
                st.error(f"❌ No se pudo buscar el archivo '{file_name}' tras varios intentos.")
                raise
        except Exception as e:
            st.error(f"Error inesperado al buscar archivo: {e}")
            raise
    
def download_file_from_drive(service, file_id, retries=3):
    """Descarga el contenido de un archivo de Drive, con reintentos."""
    for attempt in range(retries):
        try:
            request = service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            fh.seek(0)
            return fh
        except (TimeoutError, httplib2.ServerNotFoundError) as e:
            if attempt < retries - 1:
                st.toast(f"⏳ Error de red al descargar. Reintentando... ({attempt + 2}/{retries})")
                time.sleep(2 ** attempt)
            else:
                st.error(f"❌ No se pudo descargar el archivo tras varios intentos.")
                raise
        except Exception as e:
            st.error(f"Error inesperado al descargar: {e}")
            raise

# --- ESTA ES LA FUNCIÓN QUE FALTABA ---
def list_project_folders(service, root_folder_id, retries=3):
    """Lista las subcarpetas (proyectos) dentro de la carpeta raíz, con reintentos."""
    query = f"'{root_folder_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    for attempt in range(retries):
        try:
            response = service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
            return {file['name']: file['id'] for file in response.get('files', [])}
        except (TimeoutError, httplib2.ServerNotFoundError) as e:
            if attempt < retries - 1:
                st.toast(f"⏳ Error de red listando proyectos. Reintentando... ({attempt + 2}/{retries})")
                time.sleep(2 ** attempt)
            else:
                st.error("❌ No se pudieron listar los proyectos de Drive tras varios intentos.")
                return {} # Devolvemos un diccionario vacío en caso de fallo final
        except Exception as e:
            st.error(f"Error inesperado al listar proyectos: {e}")
            return {}

def sync_guiones_folders_with_index(service, project_folder_id, new_index_structure):
    """
    Compara las carpetas de guiones existentes en Drive con el nuevo índice.
    Elimina las carpetas que ya no corresponden a ningún subapartado.
    """
    st.toast("🔄 Sincronizando carpetas de guiones con el nuevo índice...")
    
    # 1. Obtener la lista de nombres de carpetas ESPERADOS según el nuevo índice
    expected_folders = set()
    if 'estructura_memoria' in new_index_structure:
        for seccion in new_index_structure.get('estructura_memoria', []):
            for subapartado_titulo in seccion.get('subapartados', []):
                # Usamos la misma lógica que en la fase 2 para crear nombres de carpeta
                nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", subapartado_titulo)
                expected_folders.add(nombre_limpio)
    
    if not expected_folders:
        st.warning("El nuevo índice no contiene subapartados. No se realizó ninguna limpieza.")
        return 0

    # 2. Obtener la lista de carpetas de guiones EXISTENTES en Google Drive
    guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
    # Reutilizamos la función que lista subcarpetas
    existing_folders_map = list_project_folders(service, guiones_main_folder_id) # Devuelve {name: id}

    # 3. Comparar y eliminar las obsoletas
    deleted_count = 0
    folders_to_delete = []
    for folder_name, folder_id in existing_folders_map.items():
        if folder_name not in expected_folders:
            folders_to_delete.append((folder_name, folder_id))

    if not folders_to_delete:
        st.toast("✅ Las carpetas de guiones ya estaban sincronizadas.")
        return 0
        
    with st.spinner(f"Eliminando {len(folders_to_delete)} carpetas de guiones obsoletas..."):
        for folder_name, folder_id in folders_to_delete:
            if delete_file_from_drive(service, folder_id):
                st.toast(f"🗑️ Carpeta obsoleta eliminada: '{folder_name}'")
                deleted_count += 1
            else:
                st.warning(f"No se pudo eliminar la carpeta obsoleta: '{folder_name}'")
    
    return deleted_count
# --- PROMPTS DE LA IA ---
# He copiado tus prompts directamente desde tu código de Colab.
PROMPT_CONSULTOR_REVISION = """
Actúas como un Consultor de Licitaciones Senior y redactor técnico experto, el mejor del mercado. Tu tarea es analizar el feedback de un cliente sobre un borrador y generar una versión mejorada que no solo corrija, sino que también proponga soluciones de alto valor.

Te proporcionaré TRES elementos:
1.  **BORRADOR ORIGINAL:** La primera versión del guion.
2.  **FEEDBACK DEL CLIENTE:** El texto del mismo documento, pero con las correcciones, ediciones o comentarios del cliente.
3.  **CONTEXTO DE LA LICITACIÓN:** Los pliegos originales para asegurar la coherencia estratégica.

Tu misión es generar una **NUEVA VERSIÓN ESTRATÉGICAMENTE SUPERIOR** del texto en formato Markdown.

## REGLAS DE ORO PARA LA REVISIÓN:
1.  **INCORPORA CORRECCIONES DIRECTAS:** Si el cliente corrige un dato o una frase, aplica ese cambio directamente. Su palabra es ley en cuanto a hechos o estilo.
2.  **SÉ UN CONSULTOR PROACTIVO (¡CLAVE!):** Si el cliente expresa una duda o un descontento (ej: "la metodología Scrum no me gusta" o "¿podemos enfocar esto de otra manera?"), NO te limites a eliminar lo antiguo. DEBES:
    a) **Analizar el problema:** Entiende por qué no le gusta la propuesta actual.
    b) **Proponer una alternativa mejor:** Basándote en tu conocimiento como licitador senior y en los pliegos, sugiere una nueva metodología, un enfoque diferente o una solución alternativa que sea más potente y tenga más probabilidades de ganar.
    c) **Justificar tu propuesta:** Explica brevemente por qué tu nueva propuesta es mejor en el contexto de esta licitación.
3.  **MANTÉN LO QUE FUNCIONA:** Conserva intactas las partes del borrador original que no recibieron feedback negativo.
4.  **FUSIÓN INTELIGENTE:** Integra todos los cambios (tanto las correcciones directas como tus nuevas propuestas) de forma natural y coherente, manteniendo el tono profesional y las reglas de oro de la redacción original.
5.  **RESPUESTA DIRECTA Y LIMPIA:** Genera únicamente el texto mejorado en Markdown. No expliques los cambios que has hecho ni uses frases introductorias.

## EJEMPLO DE ACTUACIÓN:
-   **Feedback del cliente:** "En la sección de metodología, no me convence Scrum para este proyecto, es demasiado rígido. Proponme otra cosa."
-   **Tu acción:** No solo borras Scrum. Lo reemplazas con una sección detallada sobre Kanban o Lean, explicando por qué es más flexible y adecuado para los objetivos descritos en los pliegos.

Tu objetivo final es que el cliente, al leer la nueva versión, piense: "No solo ha hecho lo que le he pedido, sino que me ha dado una solución mejor en la que no había pensado".
"""

PROMPT_PLANTILLA = """
Eres un analista de documentos extremadamente preciso.
Te daré el texto de una plantilla de memoria técnica y los Pliegos correspondientes.
Tu única tarea es convertirlo a un objeto JSON que contenga la estructura del indice y unas indicaciones para que la persona
que va a redactar la memoria técnica sepa todo lo necesario para poder redactar la memoria técnica con mayor puntuación.

## REGLAS ESTRICTAS:
1.  La estructura del documento debes sacarlo de la plantilla y las indicaciones mezclando esa información con la de los pliegos.
2.  El objeto JSON DEBE contener dos claves de nivel superior y solo dos: "estructura_memoria" y "matices_desarrollo".
3.  Para CADA apartado y subapartado, DEBES anteponer su numeración correspondiente (ej: "1. Título", "1.1. Subtítulo").
    ESTO ES OBLIGATORIO Y DEBE SER EN NÚMEROS NORMALES (1,2,3...) NADA DE LETRAS NI COSAS RARAS.
4.  La clave "estructura_memoria" contiene la lista de apartados y subapartados como un ÍNDICE.
    La lista "subapartados" SOLO debe contener los TÍTULOS numerados, NUNCA el texto de las instrucciones.
5.  Debes coger exactamente el mismo título del apartado o subapartado que existe en el texto de la plantilla, no lo modifiques.
    Mantenlo aunque esté en otro idioma.
6.  La clave "matices_desarrollo" desglosa CADA subapartado, asociando su título numerado con las INSTRUCCIONES completas.
    NO RESUMAS. DEBES CONTAR TODO LO QUE SEPAS DE ELLO.
    Llena estas indicaciones de mucho contexto útil para que alguien sin experiencia pueda redactar la memoria.
7.  DEBES INDICAR OBLIGATORIAMENTE LA LONGITUD DE CADA SUBAPARTADO.
    NO TE LO PUEDES INVENTAR. ESTE DATO ES CLAVE.
8.  Cada instrucción debe incluir. Si no tiene eso la instrucción no vale:
    - La longitud exacta de palabras del apartado (o aproximada según lo que se diga en los pliegos). No pongas en ningún caso
    "La longitud de este subapartado no está especificada en los documentos proporcionados", propon tú uno si no existe. Esta proposición debe
    ser coherente con el apartado que es y con lo que se valora en los pliegos.
    - Una explicación clara de lo que incluirá este apartado.
    - El objetivo de contenido para que este apartado sume a obtener la excelencia en la memoria técnica.
    - Cosas que no deben faltar en el apartado.

## MEJORAS AÑADIDAS:
- Responde SIEMPRE en formato JSON válido y bien estructurado. No incluyas texto fuera del objeto JSON.
- No inventes información: solo utiliza lo que aparezca en la plantilla o en los pliegos.
- Debes mostrar conocimiento de los pliegos, no puedes asumir que el que lee las intrucciones ya posee ese conociminento.
Debes explicar todo como si el que fuera a leer las indicaciones no supiera nada del tema y deba redactar todo el contenido.
- Mantén consistencia en la numeración (ejemplo: 1, 1.1, 1.1.1). Nunca mezcles números y letras.
- Si los pliegos mencionan tablas, gráficos o anexos obligatorios, añádelos en las indicaciones como recordatorio.
- Si hay discrepancias entre plantilla y pliego, PRIORIZA SIEMPRE lo que diga el pliego.
- Valida que cada subapartado en "estructura_memoria" tenga su correspondiente bloque en "matices_desarrollo".

## EJEMPLO DE ESTRUCTURA DE SALIDA OBLIGATORIA:
{
  "estructura_memoria": [
    {
      "apartado": "1. Análisis",
      "subapartados": ["1.1. Contexto", "1.2. DAFO"]
    }
  ],
  "matices_desarrollo": [
    {
      "apartado": "1. Análisis",
      "subapartado": "1.1. Contexto",
      "indicaciones": "El subapartado debe durar 5 páginas. Este subapartado debe describir el objeto de la contratación, que es la prestación de servicios de asesoramiento, mentoría y consultoría a personas emprendedoras autónomas en Galicia. El objetivo principal es apoyar la consolidación y crecimiento de 200 proyectos empresariales de trabajadores autónomos, a través de una red de mentores especializados, para potenciar sus competencias emprendedoras, mejorar su competitividad y reducir los riesgos. Se espera que se incluyan las dos modalidades de consultoría y mentoring: una estratégica para mejorar rendimiento y rentabilidad, y otra especializada para el desarrollo de una estrategia de expansión y escalabilidad, incluyendo un análisis competitivo y de mercado..."
    },
    {
      "apartado": "1. Análisis",
      "subapartado": "1.2. DAFO",
      "indicaciones": "El subapartado debe durar 5 páginas. Este subapartado debe conseguir mostrar ..."
    }
  ]
}
"""

PROMPT_PLIEGOS = """
Eres un consultor experto en licitaciones públicas, especializado en estructurar memorias técnicas para maximizar la puntuación. Tu conocimiento se basa ÚNICAMENTE en los archivos que te he proporcionado.

Tu misión es analizar los Pliegos (administrativos y técnicos) para diseñar un **índice jerárquico y estratégico** para la memoria técnica. Este índice debe responder perfectamente a todos los requisitos y, fundamentalmente, a los criterios de valoración.

## METODOLOGÍA DE ANÁLISIS OBLIGATORIA:
Para crear la estructura, seguirás estos pasos:
1.  **IDENTIFICAR APARTADOS PRINCIPALES:** Busca en los pliegos la sección de "CRITERIOS DE VALORACIÓN SUJETOS A JUICIO DE VALOR" (o similar). CADA UNO de estos criterios principales (ej: "Calidad de la Metodología", "Plan de Trabajo", "Equipo Adscrito") se convertirá en un **apartado de nivel superior** en tu estructura (ej: "1. Metodología Propuesta", "2. Plan de Trabajo", etc.).
2.  **AGRUPAR SUBAPARTADOS LÓGICAMENTE:** Para cada apartado principal que has identificado, busca en TODO el pliego (especialmente en el Pliego de Prescripciones Técnicas - PPT) los requisitos, detalles y especificaciones que correspondan a ese criterio. Estos detalles se convertirán en los **subapartados** (ej: "1.1. Fases de la Metodología", "1.2. Herramientas a utilizar").
3.  **GARANTIZAR COBERTURA TOTAL:** Asegúrate de que cada requisito relevante del pliego tenga su lugar en la estructura. Si un requisito no encaja claramente en un criterio de valoración, crea un apartado lógico para él (como "Mejoras Adicionales").

## REGLAS ESTRICTAS DE SALIDA:
0.  **LA JERARQUÍA ES CLAVE:** El objetivo es un índice bien estructurado con varios apartados principales (1, 2, 3...) y sus correspondientes subapartados (1.1, 1.2, 2.1...). **Está prohibido generar una estructura con un único apartado principal y una larga lista de subapartados.**
1.  **RESPUESTA EXCLUSIVAMENTE EN JSON:** Tu única salida debe ser un objeto JSON válido. No incluyas texto introductorio, explicaciones ni marcadores como ```json.
2.  **CLAVES PRINCIPALES FIJAS:** El objeto JSON DEBE contener dos claves de nivel superior y solo dos: "estructura_memoria" y "matices_desarrollo".
3.  **NUMERACIÓN JERÁRQUICA:** Para CADA apartado y subapartado, DEBES anteponer su numeración correspondiente (ej: "1. Título", "1.1. Subtítulo", "1.2. Subtítulo", "2. Otro Título"). Usa solo números, nunca letras.
4.  **TÍTULOS FIELES AL PLIEGO:** Utiliza los títulos y la terminología exactos de los Pliegos para los apartados y subapartados. Si el pliego no proporciona un título claro para un grupo de requisitos, puedes crear un título descriptivo y lógico.
5.  **CONTENIDO DE "matices_desarrollo":** Esta sección debe ser exhaustiva. Para CADA subapartado, las "indicaciones" deben incluir OBLIGATORIAMENTE:
    -   **Puntuación y Relevancia:** Menciona explícitamente cuántos puntos vale el criterio principal asociado y por qué este subapartado es crucial para obtenerlos.
    -   **Longitud Estimada:** Propón una longitud en páginas o palabras. Si el pliego no lo especifica, haz una estimación razonable basada en la importancia y puntuación del apartado. NUNCA digas que no está especificado.
    -   **Contenido Detallado:** Explica qué información específica del pliego se debe desarrollar aquí.
    -   **Objetivo Estratégico:** Describe qué se debe demostrar al evaluador para conseguir la máxima puntuación (ej: "El objetivo es demostrar un dominio completo del proceso X y cómo nuestra metodología mitiga los riesgos Y").
    -   **Elementos Clave a Incluir:** Lista de puntos, tablas, gráficos o datos que no pueden faltar.

## EJEMPLO DE ESTRUCTURA DE SALIDA OBLIGATORIA (CON BUENA JERARQUÍA):
{
  "estructura_memoria": [
    {
      "apartado": "1. Solución Técnica y Metodología",
      "subapartados": ["1.1. Metodología de Trabajo", "1.2. Plan de Trabajo", "1.3. Equipo de Trabajo"]
    },
    {
      "apartado": "2. Calidad del Servicio y Mejoras",
      "subapartados": ["2.1. Actuaciones adicionales", "2.2. Políticas empresariales"]
    }
  ],
  "matices_desarrollo": [
    {
      "apartado": "1. Solución Técnica y Metodología",
      "subapartado": "1.1. Metodología de Trabajo",
      "indicaciones": "Este subapartado es clave para el criterio 'Calidad de la Propuesta Técnica', valorado con 40 puntos. Se recomienda una extensión de 8 páginas. Aquí se debe detallar la metodología agile-scrum que se implementará, describiendo las fases del proyecto: Sprint 0 (Setup), Sprints de Desarrollo (ciclos de 2 semanas) y Sprint de Cierre. Es fundamental incluir un diagrama de flujo del proceso y explicar cómo las ceremonias (Daily, Planning, Review, Retro) garantizan la comunicación y la adaptación continua. El objetivo es demostrar que nuestra metodología es robusta, flexible y minimiza los riesgos de desviación del proyecto..."
    },
    {
      "apartado": "2. Calidad del Servicio y Mejoras",
      "subapartado": "2.1. Actuaciones adicionales",
      "indicaciones": "Este subapartado responde al criterio de 'Mejoras Propuestas', valorado con 15 puntos. Se recomienda una extensión de 3 páginas. Se debe proponer la implantación de un dashboard de seguimiento en tiempo real con PowerBI sin coste adicional para el cliente. Hay que detallar qué KPIs se mostrarán (ej: avance de tareas, presupuesto consumido, incidencias abiertas/cerradas) y qué beneficios aporta en términos de transparencia y toma de decisiones. No debe faltar una captura de pantalla de un dashboard de ejemplo..."
    }
  ]
}
"""

PROMPT_PREGUNTAS_TECNICAS = """
Actúa como un planificador de licitación. Te quieres presentar a una licitación y debes crear un documento enfocando el contenido que aparecerá en este para que tus compañeros vean tu propuesta
y la validen y complementen. Tu objetivo será crear una propuesta de contenido ganadora basándote en lo que se pide en los pliegos para que tus compañeros sólo den el ok
y se pueda mandar el contenido a un redactor para que simplemente profundice en lo que tu has planteado. Esa "mini memoria técnica" será la que se le dará a un compañaero que se dedica a redactar.

La estructura del documento será un indice pegando la estructrua simplemente que tendrá esa memoria técnica ("Estructura de la memoria técnica") y la propuesta de los apartados ("Propuesta de contenido para Nombre Licitación").
En la propuesta de contenido por apartado debes responder a dos preguntas: qué se debe incluir en este apartado y el contenido propuesto para ese apartado.
La primera pregunta debe ser un resumen de todo lo que se pide en el pliego para ese apartado. Debes detallar qué aspectos se valoran básandote en lo que se dice en el pliego administrativo, qué información se detallará en profundida en esa parte exclusivamente , cuales son los puntos generales que tocarás en este apartado, qué aspectos se valoran básandote en lo que se dice en el pliego técnico y las puntuaciones relativas a este apartado. Esto debe estar en párrafos y en bullet points.
La segunda pregunta debe ser tu propuesta de contenido para responder ese apartado. Esa propuesta debe enfocarse a explicar la propuesta que tu crees más óptima para obtener la mayor puntuación. Debes detallarla ampliamente de una manera esquemática enfocando en el contenido (no en la explicación) de eso que propones. Esa propuesta será analizada por tus compañeros para mejorar el enfoque.
Para responder a esa segunda pregunta, deberás crear preguntas que desengranen el contenido general de ese apartado en preguntas más pequeñas para que tus compañeros puedan ir ajustando y mejorando cada fase.
Por ejemplo, si se te habla de metodología: primero deberás leerte el pliego administrativo y ver que estructura debe tener una metodología y segundo leerte el pliego técnico y ver el contenido que debe tener. En ese caso localizaste (ampliando lo que se dice en los pliegios) que la metodología debe hablar sobre los principios que enmarcan esa propuesta, la teoría de la metodología, las actividades y el cronograma.
Con esos puntos localizados deberías escribir un párrafo amplio profundizando en esa primera pregunta de resumen de todo lo que se pide en el pliego para ese apartado y después escribir la desengranción de preguntas por apartado y dar una respuesta detallada sobre el contenido o el enfoque que deberá tener ese contenido para definir perfectamente la metodología final de esa memoria técnica.
Debe ser propuestas muy precisas, es decir, deben de ser textos que expliquen muy bien todas las actividades, metodologías y conceptos relacionados con el enfoque de una manera que la persona que lea este documento solo se dedique a matizar y a mejorar los contenidos.

Para cada apartado y subapartado del índice, desarrollarás el contenido siguiendo OBLIGATORIAMENTE estas 6 REGLAS DE ORO:

    1.  **TONO PROFESIONAL E IMPERSONAL:** Redacta siempre en tercera persona. Elimina CUALQUIER referencia personal (ej. "nosotros", "nuestra propuesta"). Usa formulaciones como "El servicio se articula en...", "La metodología implementada será...".

    2.  **CONCRECIÓN ABSOLUTA (EL "CÓMO"):** Cada afirmación general DEBE ser respaldada por una acción concreta, una herramienta específica (ej. CRM HubSpot for Startups, WhatsApp Business API), una métrica medible o un entregable tangible. Evita las frases vacías.

    3.  **ENFOQUE EN EL USUARIO FINAL (BUYER PERSONA):** Orienta todo el contenido a resolver los problemas del buyer persona objetivo de esa licitación. Demuestra un profundo conocimiento de su perfil, retos (burocracia, aislamiento) y objetivos (viabilidad, crecimiento).

    4.  **LONGITUD CONTROLADA POR PALABRAS:** El desarrollo completo de la "Propuesta de Contenido" debe tener una extensión total de entre 6.000 y 8.000 palabras. Distribuye el contenido de forma equilibrada entre los apartados para alcanzar este objetivo sin generar texto de relleno.

    5.  **PROPUESTA DE VALOR ESTRATÉGICA:** Enfócate en los resultados y el valor añadido. En esta memoria no busques adornar las ideas, centrate en mostrar las ideas de una manera fácil de ver y clara.

    6.  **ALINEACIÓN TOTAL CON EL PLIEGO (PPT):** La justificación de cada acción debe ser su alineación con los requisitos del Pliego y el valor que aporta para obtener la máxima puntuación.

    Para el desarrollo de cada apartado en la PARTE 2, usa este formato:
    -   **"Qué se debe incluir en este apartado (Análisis del Pliego)":** Resume los requisitos del PPT, criterios de evaluación y puntuación.
    -   **"Contenido Propuesto para el Apartado":** Aplica aquí las 6 Reglas de Oro, desarrollando la propuesta de forma concreta, estratégica y detallada.

En este documento solo deberán aparecer los apartados angulares de la propuesta. Se omitirán los de presentación, los de introducción y los que no vayan directamente asociados a definir lo principal de la licitación. Normalmente lo prinicipal es la metodología, las actividades que se van a hacer y la planificación con su cronograma correspondiente.

Te proporcionaré DOS elementos clave:
1.  El texto completo de los documentos base (Pliegos y/o plantilla).
2.  La estructura que se ha generado en el mensaje anterior con los apartados y las anotaciones.
"""

PROMPT_PREGUNTAS_TECNICAS_INDIVIDUAL = """
Actúa como un planificador de licitación. Te quieres presentar a una licitación y debes crear un documento enfocando el contenido que aparecerá en este para que tus compañeros vean tu propuesta
y la validen y complementen. Tu objetivo será crear una propuesta de contenido ganadora basándote en lo que se pide en los pliegos para que tus compañeros sólo den el ok
y se pueda mandar el contenido a un redactor para que simplemente profundice en lo que tu has planteado. Esa "mini memoria técnica" será la que se le dará a un compañaero que se dedica a redactar.

!! Tu respuesta debe centrarse EXCLUSIVAMENTE en el apartado proporcionado. No incluyas un índice general ni el título "Propuesta de contenido para...". Empieza directamente con el desarrollo del apartado. !!
Para el apartado proporcionado, debes responder a dos preguntas: "qué se debe incluir en este apartado" y "el contenido propuesto para ese apartado".

La primera pregunta ("Qué se debe incluir...") debe ser un resumen de todo lo que se pide en el pliego para ese apartado. Debes detallar qué aspectos se valoran, qué información se detallará en profundidad, cuáles son los puntos generales que tocarás, qué aspectos se valoran según el pliego técnico y las puntuaciones relativas. Usa párrafos y bullet points.

La segunda pregunta ("Contenido propuesto...") debe ser tu propuesta de contenido para obtener la mayor puntuación. Detállala ampliamente de manera esquemática, enfocando en el contenido (no en la explicación). Desgrana el contenido general en preguntas más pequeñas y da respuestas detalladas que expliquen muy bien las actividades, metodologías y conceptos.

Para cada apartado y subapartado del índice, desarrollarás el contenido siguiendo OBLIGATORIAMENTE estas 6 REGLAS DE ORO:

    1.  **TONO PROFESIONAL E IMPERSONAL:** Redacta siempre en tercera persona. Elimina CUALQUIER referencia personal (ej. "nosotros", "nuestra propuesta"). Usa formulaciones como "El servicio se articula en...", "La metodología implementada será...".

    2.  **CONCRECIÓN ABSOLUTA (EL "CÓMO"):** Cada afirmación general DEBE ser respaldada por una acción concreta, una herramienta específica (ej. CRM HubSpot for Startups, WhatsApp Business API), una métrica medible o un entregable tangible. Evita las frases vacías.

    3.  **ENFOQUE EN EL USUARIO FINAL (BUYER PERSONA):** Orienta todo el contenido a resolver los problemas del buyer persona objetivo de esa licitación. Demuestra un profundo conocimiento de su perfil, retos (burocracia, aislamiento) y objetivos (viabilidad, crecimiento).

    4.  **LONGITUD CONTROLADA POR PALABRAS:** El desarrollo completo de la "Propuesta de Contenido" debe tener una extensión total de entre 6.000 y 8.000 palabras. Distribuye el contenido de forma equilibrada entre los apartados para alcanzar este objetivo sin generar texto de relleno.

    5.  **PROPUESTA DE VALOR ESTRATÉGICA:** Enfócate en los resultados y el valor añadido. En esta memoria no busques adornar las ideas, centrate en mostrar las ideas de una manera fácil de ver y clara.

    6.  **ALINEACIÓN TOTAL CON EL PLIEGO (PPT):** La justificación de cada acción debe ser su alineación con los requisitos del Pliego y el valor que aporta para obtener la máxima puntuación.

    Para el desarrollo de cada apartado en la PARTE 2, usa este formato:
    -   **"Qué se debe incluir en este apartado (Análisis del Pliego)":** Resume los requisitos del PPT, criterios de evaluación y puntuación.
    -   **"Contenido Propuesto para el Apartado":** Aplica aquí las 6 Reglas de Oro, desarrollando la propuesta de forma concreta, estratégica y detallada.

En este documento solo deberán aparecer los apartados angulares de la propuesta. Se omitirán los de presentación, los de introducción y los que no vayan directamente asociados a definir lo principal de la licitación. Normalmente lo prinicipal es la metodología, las actividades que se van a hacer y la planificación con su cronograma correspondiente.

Te proporcionaré TRES elementos clave:
1.  El texto completo de los documentos base (Pliegos).
2.  Las indicaciones para el apartado específico que debes desarrollar (extraídas de un JSON de estructura).
3.  Documentación de apoyo adicional (opcional) que el usuario haya subido para este apartado.
"""


PROMPT_REGENERACION = """
Actúas como un editor experto que refina una estructura JSON para una memoria técnica.
Te proporcionaré TRES elementos clave:
1.  Los documentos originales (Pliegos y/o plantilla).
2.  La estructura JSON que se generó en un primer intento.
3.  Las INSTRUCCIONES DE UN USUARIO con los cambios que desea.

Tu única tarea es generar una **NUEVA VERSIÓN MEJORADA** del objeto JSON que incorpore a la perfección los cambios solicitados por el usuario.

## REGLAS OBLIGATORIAS:
-   **MANTÉN TODAS LAS REGLAS DEL PROMPT ORIGINAL:** El formato de salida debe seguir siendo un JSON válido con las claves "estructura_memoria" y "matices_desarrollo", la numeración debe ser correcta (1, 1.1, etc.), y las indicaciones deben ser detalladas.
-   **INCORPORA EL FEEDBACK:** Lee atentamente las instrucciones del usuario y aplícalas a la nueva estructura. Por ejemplo, si pide "une los apartados 1.1 y 1.2", debes hacerlo. Si pide "el apartado 2 debe hablar sobre la experiencia del equipo", debes modificar las indicaciones de ese apartado.
-   **NO PIERDAS INFORMACIÓN:** Si el usuario solo pide cambiar el apartado 3, los apartados 1, 2, 4, etc., deben permanecer intactos en la nueva versión.
-   **SÉ PRECISO:** No inventes nuevos apartados a menos que el usuario te lo pida explícitamente. Céntrate únicamente en aplicar las correcciones solicitadas.

Genera únicamente el objeto JSON corregido. No incluyas ningún texto fuera de él.
"""

PROMPT_DESARROLLO = """
**SYSTEM DIRECTIVES: NON-NEGOTIABLE RULES FOR OUTPUT GENERATION.**
**FAILURE TO FOLLOW THESE RULES WILL INVALIDATE THE ENTIRE RESPONSE.**

1.  **FORBIDDEN CONTENT:** You are STRICTLY PROHIBITED from generating ANY of the following:
    *   **Meta-commentary or self-analysis.** This includes explaining the code you are about to generate or have generated. **Forbidden phrases include, but are not limited to:** "Este código crea un diagrama...", "Recuerda guardar este código...", "This HTML creates...", etc.
    *   **Mermaid code** (e.g., `mermaid graph LR...`).
    *   **Standalone CSS** (e.g., `<style>...</style>`). All styling is embedded in the templates.
    *   **Instructions, placeholders, or comments to the user** (e.g., "(Insertar aquí...)", "[Complete here]"). Generate the final content yourself.

2.  **ALLOWED CONTENT:** Your entire response MUST consist ONLY of:
    *   **Spanish (castellano) Markdown text.**
    *   **OR** one of the two complete HTML templates provided below for visual elements. NO other formats are permitted.

3.  **OUTPUT STRUCTURE:**
    *   Your response must start directly with the first paragraph or the `<!DOCTYPE html>` tag. DO NOT repeat the subsection title.
    *   Your response MUST end immediately after the final Markdown text or the closing `</html>` tag. DO NOT add any extra text, newlines, or explanations after the content.

---
## YOUR PERSONA AND TASK

You are an expert consultant drafting a technical proposal for a public tender. Your task is to generate the complete content for a specific subsection. The writing must be professional, direct, and strictly adhere to the System Directives above.

---
## VISUAL TOOLS: THE ONLY TWO ALLOWED HTML TEMPLATES

**OPTION A: SIMPLE LIST TEMPLATE (For benefits, pillars, features)**
```html
<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale-1.0"><title>Visual Element</title><style>@import url('https://fonts.googleapis.com/css2?family=Urbanist:wght@400;600;700&display=swap');body {{{{ font-family: 'Urbanist', sans-serif; background-color: #f0f2f5; display: flex; justify-content: center; align-items: center; padding: 20px; width: 800px; box-sizing: border-box; }}}} .card {{{{ background-color: white; border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); padding: 25px; width: 100%; max-width: 750px; border-top: 5px solid #0046C6; }}}} h2 {{{{ color: #0046C6; text-align: center; margin-top: 0; font-size: 24px; font-weight: 700; }}}} ul {{{{ list-style-type: none; padding: 0; }}}} li {{{{ display: flex; align-items: center; margin-bottom: 15px; font-size: 16px; color: #333; }}}} li::before {{{{ content: '✔'; color: #32CFAA; font-size: 20px; font-weight: bold; margin-right: 15px; }}}}</style></head><body><div class="card"><h2><!-- TÍTULO AQUÍ --></h2><ul><!-- LISTA DE ELEMENTOS AQUÍ --></ul></div></body></html>
**OPTION B: MULTI-COLUMN INFOGRAPHIC TEMPLATE (For phases, flowcharts, action areas)**
<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale-1.0"><title>Infographic Element</title><style>@import url('https://fonts.googleapis.com/css2?family=Urbanist:wght@400;600;700&display=swap');body {{{{ font-family: 'Urbanist', sans-serif; display: flex; justify-content: center; align-items: center; padding: 20px; background-color: #f8f9fa; width: 800px; box-sizing: border-box; }}}} .container {{{{ background-color: #ffffff; border: 2px dashed #e0e0e0; border-radius: 15px; padding: 25px; width: 100%; box-shadow: 0 4px 12px rgba(0,0,0,0.05); }}}} h2 {{{{ color: #0046C6; text-align: center; font-size: 26px; font-weight: 700; margin-bottom: 25px; letter-spacing: -0.5px; }}}} .columns {{{{ display: flex; justify-content: space-around; gap: 20px; }}}} .column {{{{ flex: 1; text-align: center; padding: 15px; border-top: 4px solid; border-radius: 8px; background-color: #fdfdfd; }}}} .column-icon {{{{ width: 30px; height: 30px; border-radius: 50%; margin: 0 auto 15px auto; }}}} .column h3 {{{{ font-size: 16px; font-weight: 600; color: #333; margin-bottom: 10px; }}}} .column ul {{{{ list-style: none; padding: 0; margin: 0; text-align: left; }}}} .column li {{{{ font-size: 13px; color: #555; margin-bottom: 8px; line-height: 1.5; }}}} .color-1 {{{{ border-color: #FBC02D; }}}} .icon-1 {{{{ background-color: #FBC02D; }}}} .color-2 {{{{ border-color: #4CAF50; }}}} .icon-2 {{{{ background-color: #4CAF50; }}}} .color-3 {{{{ border-color: #90CAF9; }}}} .icon-3 {{{{ background-color: #90CAF9; }}}} .color-4 {{{{ border-color: #F44336; }}}} .icon-4 {{{{ background-color: #F44336; }}}}</style></head><body><div class="container"><h2><!-- TÍTULO --></h2><div class="columns"><div class="column color-1"><div class="column-icon icon-1"></div><h3><!-- Título Col 1 --></h3><ul><li><!-- Punto 1 --></li></ul></div><div class="column color-2"><div class="column-icon icon-2"></div><h3><!-- Título Col 2 --></h3><ul><li><!-- Punto 1 --></li></ul></div><div class="column color-3"><div class="column-icon icon-3"></div><h3><!-- Título Col 3 --></h3><ul><li><!-- Punto 1 --></li></ul></div></div></div></body></html>
**YOUR SPECIFIC TASK (in English, for clarity to the model)**
You must now generate a JSON object for the following subsection.
Main Section: "{apartado_titulo}"
Subsection to Draft: "{subapartado_titulo}"
Guidelines for this Subsection: "{indicaciones}"
FINAL OUTPUT FORMAT (STRICT):
Your response MUST be ONLY a single, valid JSON object (no ```json fences). It must contain a single key "plan_de_prompts" which is a list of objects. Each object must follow this exact structure:
{{{{
"apartado_referencia": "{apartado_titulo}",
"subapartado_referencia": "{subapartado_titulo}",
"prompt_id": "A unique ID (e.g., PROMPT_2_1_A). Add '_HTML_VISUAL' if it's an HTML element.",
"prompt_para_asistente": "The specific and detailed prompt for the assistant to generate the Spanish content, following all system directives."
}}}}
"""

# =============================================================================
#           NUEVOS PROMPTS: AÑADE ESTOS A TU SCRIPT
# =============================================================================

PROMPT_GENERAR_INTRODUCCION = """
Actúas como un estratega experto en la redacción de propuestas de licitación. Tu tarea es escribir un apartado de **Introducción** conciso y persuasivo, basándote en el contenido completo de la memoria técnica que te proporcionaré.

## REGLAS ESTRICTAS:
1.  **ENFOQUE EN LA SOLUCIÓN:** No te limites a describir el documento ("En esta memoria se describirá..."). En su lugar, resume la **propuesta de valor** y la solución que se ofrece. Empieza con fuerza.
2.  **SÍNTESIS ESTRATÉGICA:** Lee y comprende la totalidad del documento para identificar los puntos más fuertes de la propuesta (ej: una metodología innovadora, un equipo experto, mejoras significativas) y destácalos brevemente.
3.  **ESTRUCTURA DEL CONTENIDO:** Tras presentar la propuesta de valor, esboza de forma narrativa la estructura del documento, guiando al lector sobre lo que encontrará. (ej: "A lo largo de los siguientes apartados, se detallará la metodología de trabajo propuesta, seguida de un exhaustivo plan de trabajo y la presentación del equipo técnico adscrito al proyecto, finalizando con las mejoras adicionales que aportan un valor diferencial.").
4.  **TONO PROFESIONAL:** Mantén un tono formal, seguro y orientado a resultados.
5.  **SALIDA DIRECTA:** Genera únicamente el texto de la introducción en formato Markdown. No incluyas el título "Introducción" ni ningún otro comentario.

**Ejemplo de inicio:** "El presente proyecto aborda la necesidad de [problema principal del cliente] a través de una solución integral que combina [pilar 1 de la solución] con [pilar 2 de la solución], garantizando [resultado clave para el cliente]."
"""

PROMPT_COHESION_FINAL = """
Actúas como un Redactor Técnico Senior y Editor experto, con la misión de transformar un borrador bien estructurado en un documento final impecable, coherente y persuasivo. Te proporcionaré el texto completo de una memoria técnica (excluyendo la introducción).

Tu tarea es reescribir y mejorar el texto aplicando las siguientes directrices de forma OBLIGATORIA:

1.  **GARANTIZAR LA COHESIÓN GLOBAL (MÁXIMA PRIORIDAD):**
    *   **Crea Puentes Lógicos:** Revisa las transiciones entre apartados y subapartados. Asegúrate de que fluyan de manera natural. Añade frases o párrafos cortos de transición si es necesario (ej: "Partiendo de la metodología Agile-Scrum descrita anteriormente, el plan de trabajo se estructura en...", "Para dar soporte a estas actividades, se utilizará el equipo técnico que se detalla a continuación...").
    *   **Consistencia Terminológica:** Unifica la terminología. Si en un apartado se habla de "Development Team" y en otro de "Equipo de Desarrollo", elige el término más adecuado (preferiblemente en español) y úsalo de forma consistente en todo el documento.
    *   **REFERENCIAS CRUZADAS (REQUISITO CLAVE):** Tu principal valor es conectar las ideas a lo largo del documento. Identifica activamente oportunidades para referenciar información ya mencionada. Si en el apartado 5 se retoma un concepto del apartado 1, debes indicarlo explícitamente. **Usa frases como: "Como se detalló en el apartado 1.1, la metodología Agile-Scrum...", "Este enfoque mitiga los riesgos identificados previamente en el análisis DAFO (ver apartado 2.3).", o "El equipo responsable de esta tarea es el presentado en la sección de Equipo Adscrito."**

2.  **MEJORAR EL ESTILO Y LA LEGIBILIDAD:**
    *   **Voz Activa y Directa:** Transforma frases pasivas en activas para dar más fuerza al texto.
    *   **Claridad y Concisión:** Elimina redundancias, jerga innecesaria y frases de relleno.
    *   **Formato y Estilo:** Asegúrate de que todos los títulos y subtítulos sigan un formato consistente de "Sentence case" (solo la primera letra del título en mayúscula, a menos que sea un nombre propio). Ej: "Diagrama de flujo del proceso" en lugar de "Diagrama de Flujo del Proceso".

3.  **REGLAS DE EXCLUSIÓN (LO QUE NO DEBES HACER):**
    *   **No Alterar Datos Clave:** No cambies datos técnicos, nombres de herramientas, cifras, duraciones de Sprints o cualquier información factual del borrador original. Tu trabajo es de estilo y cohesión, no de contenido.
    *   **No Añadir Nueva Información:** No inventes nuevas funcionalidades o fases que no estuvieran en el texto original.
    *   **No Incluir Meta-Comentarios:** Tu salida debe ser ÚNICAMENTE el texto final y pulido en formato Markdown. No expliques los cambios que has hecho.

El objetivo final es que el documento se lea como si hubiera sido escrito por una única persona experta, con una narrativa fluida y un argumento comercial sólido y cohesionado de principio a fin.
"""
# =============================================================================
#              NUEVAS FUNCIONES: AUTENTICACIÓN Y GOOGLE DRIVE
# =============================================================================

def get_google_flow():
    """Crea y devuelve el objeto Flow de Google OAuth."""
    return Flow.from_client_config(
        client_config=CLIENT_CONFIG,
        scopes=SCOPES,
        redirect_uri=st.secrets["GOOGLE_REDIRECT_URI"]
    )

def get_credentials():
    """Obtiene las credenciales, forzando un nuevo login si los scopes han cambiado."""
    
    # --- !! ESTA ES LA NUEVA LÓGICA DE COMPROBACIÓN !! ---
    if 'credentials' in st.session_state and st.session_state.credentials:
        creds = st.session_state.credentials
        # Comprobamos si todos los scopes que necesitamos están en las credenciales guardadas.
        if not all(scope in creds.scopes for scope in SCOPES):
            # Si faltan scopes, las credenciales no son válidas. Las borramos.
            del st.session_state.credentials
            # Forzamos al usuario a la página de inicio para que se loguee de nuevo.
            go_to_landing()
            st.rerun() # Detenemos la ejecución actual y recargamos
    # --- !! FIN DE LA NUEVA LÓGICA !! ---

    if 'credentials' in st.session_state and st.session_state.credentials:
        creds = st.session_state.credentials
        if creds.expired and creds.refresh_token:
            try:
                # Necesitamos importar Request de google.auth.transport.requests
                from google.auth.transport.requests import Request
                creds.refresh(Request())
                st.session_state.credentials = creds
            except Exception as e:
                # Si el refresh token falla (p.ej. ha sido revocado), borramos credenciales
                del st.session_state.credentials
                go_to_landing()
                st.rerun()
        return creds

    if 'code' in st.query_params:
        try:
            flow = get_google_flow()
            flow.fetch_token(code=st.query_params['code'])
            st.session_state.credentials = flow.credentials
            # Limpiamos los parámetros de la URL para evitar bucles
            st.query_params.clear()
            st.rerun()
        except Exception as e:
            # Si hay un error (como 'scope has changed'), lo mostramos y limpiamos el estado.
            st.error(f"Error al obtener el token: {e}")
            if 'credentials' in st.session_state:
                del st.session_state.credentials
            # Damos la opción de reintentar
            st.button("Reintentar inicio de sesión")
            st.stop() # Detenemos la ejecución para que no continúe con error
            
    return None
    
def build_drive_service(credentials):
    """Construye y devuelve el objeto de servicio de la API de Drive."""
    try:
        return build('drive', 'v3', credentials=credentials)
    except HttpError as error:
        st.error(f"No se pudo crear el servicio de Drive: {error}")
        return None

def find_or_create_folder(service, folder_name, parent_id=None):
    """Busca una carpeta por nombre. Si no la encuentra, la crea."""
    query = f"name = '{folder_name}' and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    
    response = service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
    files = response.get('files', [])
    
    if files:
        return files[0]['id']
    else:
        file_metadata = {'name': folder_name, 'mimeType': 'application/vnd.google-apps.folder'}
        if parent_id:
            file_metadata['parents'] = [parent_id]
        
        folder = service.files().create(body=file_metadata, fields='id').execute()
        st.toast(f"Carpeta '{folder_name}' creada en tu Drive.")
        return folder.get('id')

def list_project_folders(service, root_folder_id):
    """Lista las subcarpetas (proyectos) dentro de la carpeta raíz."""
    query = f"'{root_folder_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
    response = service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
    return {file['name']: file['id'] for file in response.get('files', [])}
    
def get_files_in_project(service, project_folder_id):
    """Obtiene los archivos dentro de una carpeta de proyecto."""
    query = f"'{project_folder_id}' in parents and trashed = false"
    response = service.files().list(q=query, spaces='drive', fields='files(id, name, mimeType)').execute()
    return response.get('files', [])
    
def download_file_from_drive(service, file_id):
    """Descarga el contenido de un archivo de Drive y lo devuelve como BytesIO."""
    request = service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
    fh.seek(0)
    return fh
    
# --- FUNCIONES AUXILIARES DE BACKEND ---
def limpiar_respuesta_json(texto_sucio):
    if not isinstance(texto_sucio, str): return ""
    match_bloque = re.search(r'```(?:json)?\s*(\{.*\})\s*```', texto_sucio, re.DOTALL)
    if match_bloque: return match_bloque.group(1).strip()
    match_objeto = re.search(r'\{.*\}', texto_sucio, re.DOTALL)
    if match_objeto: return match_objeto.group(0).strip()
    return ""

def agregar_markdown_a_word(documento, texto_markdown):
    patron_encabezado = re.compile(r'^(#+)\s+(.*)')
    patron_lista_numerada = re.compile(r'^\s*\d+\.\s+')
    patron_lista_viñeta = re.compile(r'^\s*[\*\-]\s+')
    def procesar_linea_con_negritas(parrafo, texto):
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'): parrafo.add_run(parte[2:-2]).bold = True
            elif parte: parrafo.add_run(parte)
    for linea in texto_markdown.split('\n'):
        linea_limpia = linea.strip()
        if not linea_limpia: continue
        match_encabezado = patron_encabezado.match(linea_limpia)
        if match_encabezado:
            documento.add_heading(match_encabezado.group(2).strip(), level=min(len(match_encabezado.group(1)), 4))
            continue
        if patron_lista_numerada.match(linea_limpia):
            p = documento.add_paragraph(style='List Number')
            procesar_linea_con_negritas(p, patron_lista_numerada.sub('', linea_limpia))
        elif patron_lista_viñeta.match(linea_limpia):
            p = documento.add_paragraph(style='List Bullet')
            procesar_linea_con_negritas(p, patron_lista_viñeta.sub('', linea_limpia))
        else:
            p = documento.add_paragraph()
            procesar_linea_con_negritas(p, linea_limpia)

def mostrar_indice_desplegable(estructura_memoria):
    if not estructura_memoria:
        st.warning("No se encontró una estructura de memoria para mostrar.")
        return
    st.subheader("Índice Propuesto")
    for seccion in estructura_memoria:
        apartado_titulo = seccion.get("apartado", "Apartado sin título")
        subapartados = seccion.get("subapartados", [])
        with st.expander(f"**{apartado_titulo}**"):
            if subapartados:
                for sub in subapartados: st.markdown(f"- {sub}")
            else: st.markdown("_Este apartado no tiene subapartados definidos._")

def sanitize_json_string(json_str):
    """
    Elimina caracteres de control inválidos de un string antes de parsearlo como JSON.
    Estos caracteres a veces son introducidos por el LLM al procesar PDFs/DOCX.
    """
    # Expresión regular para encontrar caracteres de control ASCII (0-31),
    # excepto los que son válidos en JSON strings si están escapados (tab, newline, etc.).
    # Esta regex busca los que causan errores de parseo.
    control_chars_regex = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')
    
    # Reemplazamos los caracteres problemáticos por una cadena vacía.
    sanitized_str = control_chars_regex.sub('', json_str)
    return sanitized_str
def wrap_html_fragment(html_fragment):
    """
    Toma un fragmento de HTML (ej: un <table> o <div>) y lo envuelve en una
    estructura HTML completa con los estilos CSS necesarios para renderizarlo.
    """
    # Si ya es un documento completo, lo devuelve tal cual.
    if html_fragment.strip().startswith('<!DOCTYPE html>'):
        return html_fragment

    # Estilos CSS extraídos de tu PROMPT_DESARROLLO.
    # Son necesarios para que las tablas y cards se vean bien.
    css_styles = """
        @import url('https://fonts.googleapis.com/css2?family=Urbanist:wght@400;600;700&display=swap');
        body {
            font-family: 'Urbanist', sans-serif;
            background-color: #f0f2f5;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
            width: 800px;
            box-sizing: border-box;
        }
        .card {
            background-color: white;
            border-radius: 10px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
            padding: 25px;
            width: 100%;
            max-width: 750px;
            border-top: 5px solid #0046C6;
        }
        h2 {
            color: #0046C6;
            text-align: center;
            margin-top: 0;
            font-size: 24px;
            font-weight: 700;
        }
        ul { list-style-type: none; padding: 0; }
        li {
            display: flex;
            align-items: center;
            margin-bottom: 15px;
            font-size: 16px;
            color: #333;
        }
        li::before {
            content: '✔';
            color: #32CFAA;
            font-size: 20px;
            font-weight: bold;
            margin-right: 15px;
        }
        /* Estilos adicionales para tablas */
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
            font-size: 15px;
        }
        th, td {
            padding: 12px 15px;
            border: 1px solid #ddd;
            text-align: left;
        }
        th {
            background-color: #f5f5f5;
            font-weight: 600;
            color: #333;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
    """
    
    # Plantilla HTML completa
    full_html_template = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Visual Element</title>
        <style>{css_styles}</style>
    </head>
    <body>
        {html_fragment}
    </body>
    </html>
    """
    return full_html_template

def html_a_imagen(html_string, output_filename="temp_image.png"):
    """
    Convierte una cadena de HTML en una imagen PNG, encontrando automáticamente
    el ejecutable wkhtmltoimage en el entorno de Streamlit Cloud.
    """
    try:
        # En Streamlit Cloud, el ejecutable se instala en una ruta accesible.
        # 'which' es un comando de Linux para encontrar la ruta de un programa.
        path_wkhtmltoimage = os.popen('which wkhtmltoimage').read().strip()

        if not path_wkhtmltoimage:
            st.error("❌ El ejecutable 'wkhtmltoimage' no se encontró. Asegúrate de que 'wkhtmltopdf' está en tu packages.txt y que la app ha sido reiniciada.")
            return None

        # Crea una configuración para imgkit apuntando al ejecutable encontrado.
        config = imgkit.config(wkhtmltoimage=path_wkhtmltoimage)
        
        # Opciones para mejorar la calidad y el tamaño de la imagen
        options = {
            'format': 'png',
            'encoding': "UTF-8",
            'width': '800',  # Un ancho fijo para consistencia
            'quiet': ''      # Suprime la salida de la consola
        }

        # Genera la imagen desde la cadena de HTML
        imgkit.from_string(html_string, output_filename, config=config, options=options)
        
        if os.path.exists(output_filename):
            return output_filename
        else:
            st.warning(f"imgkit ejecutado pero el archivo '{output_filename}' no fue creado.")
            return None

    except Exception as e:
        st.error(f"Ocurrió un error al convertir HTML a imagen: {e}")
        st.code(f"Path de wkhtmltoimage intentado: {os.popen('which wkhtmltoimage').read().strip()}", language="bash")
        return None
# AÑADE ESTA NUEVA FUNCIÓN A TU SCRIPT
def limpiar_respuesta_narrativa(texto_ia):
    """
    Limpia la respuesta de la IA para la redacción final, eliminando:
    - Bloques de código JSON.
    - Frases introductorias comunes.
    - El propio título del subapartado si la IA lo repite.
    """
    if not isinstance(texto_ia, str):
        return ""

    # Eliminar bloques de código JSON completos
    texto_limpio = re.sub(r'```json\s*\{.*?\}\s*```', '', texto_ia, flags=re.DOTALL)
    
    # Eliminar frases introductorias comunes (puedes añadir más)
    frases_a_eliminar = [
        r'^\s*Aquí tienes el contenido para el subapartado.*?:',
        r'^\s*Claro, aquí está la redacción para.*?:',
        r'^\s*A continuación se presenta el contenido detallado:',
        r'^\s*##\s*.*?$' # Elimina cualquier título Markdown que la IA pueda añadir
    ]
    for patron in frases_a_eliminar:
        texto_limpio = re.sub(patron, '', texto_limpio, flags=re.IGNORECASE | re.MULTILINE).strip()

    return texto_limpio

# --- NAVEGACIÓN Y GESTIÓN DE ESTADO (actualizada) ---
if 'page' not in st.session_state: st.session_state.page = 'landing'
if 'credentials' not in st.session_state: st.session_state.credentials = None
if 'drive_service' not in st.session_state: st.session_state.drive_service = None
if 'selected_project' not in st.session_state: st.session_state.selected_project = None

def go_to_project_selection(): st.session_state.page = 'project_selection'
def go_to_landing(): st.session_state.page = 'landing'
def go_to_phase1(): st.session_state.page = 'phase_1'
def go_to_phase1_results(): st.session_state.page = 'phase_1_results'
def go_to_phase2():
    st.session_state.page = 'phase_2'
def go_to_phase3(): st.session_state.page = 'phase_3'
def go_to_phase4(): st.session_state.page = 'phase_4' # <-- AÑADE ESTA LÍNEA

def back_to_project_selection_and_cleanup():
    for key in ['generated_structure', 'word_file', 'uploaded_template', 'uploaded_pliegos', 'selected_project']:
        if key in st.session_state: del st.session_state[key]
    go_to_project_selection()

def handle_full_regeneration(model):
    """
    Función centralizada que genera un índice completamente nuevo desde cero
    analizando los archivos de la carpeta 'Pliegos'.
    """
    if not st.session_state.get('drive_service') or not st.session_state.get('selected_project'):
        st.error("Error de sesión. No se puede iniciar la regeneración.")
        return False

    with st.spinner("Descargando archivos de 'Pliegos' y re-analizando desde cero..."):
        try:
            service = st.session_state.drive_service
            project_folder_id = st.session_state.selected_project['id']
            pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
            document_files = get_files_in_project(service, pliegos_folder_id)

            if not document_files:
                st.warning("No se encontraron archivos en la carpeta 'Pliegos' para analizar.")
                return False

            downloaded_files_for_ia = []
            for file in document_files:
                file_content_bytes = download_file_from_drive(service, file['id'])
                downloaded_files_for_ia.append({"mime_type": file['mimeType'], "data": file_content_bytes.getvalue()})

            contenido_ia = [PROMPT_PLIEGOS]
            contenido_ia.extend(downloaded_files_for_ia)
            
            generation_config = genai.GenerationConfig(response_mime_type="application/json")
            response = model.generate_content(contenido_ia, generation_config=generation_config)
            
            json_limpio_str = limpiar_respuesta_json(response.text)
            if json_limpio_str:
                informacion_estructurada = json.loads(json_limpio_str)
                st.session_state.generated_structure = informacion_estructurada
                st.session_state.uploaded_pliegos = document_files # Actualizamos la lista de pliegos por si acaso
                st.toast("✅ ¡Índice regenerado desde cero con éxito!")
                return True
            else:
                st.error("La IA devolvió una respuesta vacía o no válida.")
                return False
        except Exception as e:
            st.error(f"Ocurrió un error durante la regeneración completa: {e}")
            return False
# =============================================================================
#                 PÁGINAS DE LA APLICACIÓN (NUEVA VERSIÓN)
# =============================================================================

def landing_page():
    """Pantalla de bienvenida que ahora incluye el inicio de sesión con Google."""
    col1, col_center, col3 = st.columns([1, 2, 1])
    with col_center:
        st.write("")
        st.markdown(f'<div style="text-align: center;"><img src="https://raw.githubusercontent.com/soporte2-tech/appfront/main/imagen.png" width="150"></div>', unsafe_allow_html=True)
        st.write("")
        st.markdown("<div style='text-align: center;'><h1>Asistente Inteligente para Memorias Técnicas</h1></div>", unsafe_allow_html=True)
        st.markdown("<div style='text-align: center;'><h3>Optimiza y acelera la creación de tus propuestas de licitación</h3></div>", unsafe_allow_html=True)
        st.markdown("---")
        st.info("Para empezar, necesitas dar permiso a la aplicación para que gestione los proyectos en tu Google Drive.")
        
        # Generamos la URL de autenticación
        flow = get_google_flow()
        auth_url, _ = flow.authorization_url(prompt='consent')
        
        # Usamos st.link_button para una experiencia de usuario limpia
        st.link_button("🔗 Conectar con Google Drive", auth_url, use_container_width=True, type="primary")

def project_selection_page():
    """Nueva página para seleccionar o crear un proyecto en Google Drive."""
    st.markdown("<h3>Selección de Proyecto</h3>", unsafe_allow_html=True)
    st.markdown("Elige un proyecto existente de tu Google Drive o crea uno nuevo para empezar.")
    st.markdown("---")
    
    # Construimos el servicio de Drive si no existe
    if 'drive_service' not in st.session_state or not st.session_state.drive_service:
        st.session_state.drive_service = build_drive_service(st.session_state.credentials)
    
    service = st.session_state.drive_service
    # Manejo de error si el servicio no se puede crear
    if not service:
        st.error("No se pudo conectar con Google Drive. Por favor, intenta volver a la página de inicio y reconectar.")
        if st.button("← Volver a Inicio"):
            # Limpiamos las credenciales para forzar un nuevo login
            for key in ['credentials', 'drive_service']:
                if key in st.session_state:
                    del st.session_state[key]
            go_to_landing()
            st.rerun()
        return

    # Buscamos o creamos la carpeta raíz y listamos los proyectos
    with st.spinner("Accediendo a tu Google Drive..."):
        root_folder_id = find_or_create_folder(service, ROOT_FOLDER_NAME)
        projects = list_project_folders(service, root_folder_id)
    
    with st.container(border=True):
        st.subheader("1. Elige un proyecto existente")
        if not projects:
            st.info("Aún no tienes proyectos. Crea uno nuevo en el paso 2.")
        else:
            # Añadimos una opción vacía para que el usuario deba elegir activamente
            project_names = ["-- Selecciona un proyecto --"] + list(projects.keys())
            selected_name = st.selectbox("Selecciona tu proyecto:", project_names)
            
            if st.button("Cargar Proyecto Seleccionado", type="primary"):
                if selected_name != "-- Selecciona un proyecto --":
                    st.session_state.selected_project = {"name": selected_name, "id": projects[selected_name]}
                    st.toast(f"Proyecto '{selected_name}' cargado.")
                    go_to_phase1()
                    st.rerun()
                else:
                    st.warning("Por favor, selecciona un proyecto de la lista.")

    with st.container(border=True):
        st.subheader("2. O crea un nuevo proyecto")
        new_project_name = st.text_input("Nombre del nuevo proyecto (ej: Licitación Metro Madrid 2024)", key="new_project_name_input")
        if st.button("Crear y Empezar Nuevo Proyecto"):
            if not new_project_name.strip():
                st.warning("Por favor, introduce un nombre para el proyecto.")
            elif new_project_name in projects:
                st.error("Ya existe un proyecto con ese nombre. Por favor, elige otro.")
            else:
                with st.spinner(f"Creando carpeta '{new_project_name}' en tu Drive..."):
                    new_project_id = find_or_create_folder(service, new_project_name, parent_id=root_folder_id)
                    st.session_state.selected_project = {"name": new_project_name, "id": new_project_id}
                    st.success(f"¡Proyecto '{new_project_name}' creado! Ya puedes cargar los documentos.")
                    go_to_phase1()
                    st.rerun()
# AÑADE ESTA FUNCIÓN A TUS FUNCIONES AUXILIARES (Y BORRA LA ANTIGUA)

# REEMPLAZA TU FUNCIÓN de limpieza con esta versión más potente

def limpiar_respuesta_final(texto_ia):
    """
    Limpia de forma agresiva la respuesta de la IA, eliminando todo
    el "meta-texto", explicaciones, y bloques de código mal formateados.
    """
    if not isinstance(texto_ia, str):
        return ""

    # <-- ¡NUEVA REGLA! Elimina comentarios específicos sobre la creación de diagramas/código.
    # Esto busca frases que empiezan con "Este código..." y terminan con "...visualizar el diagrama." y lo elimina todo.
    texto_limpio = re.sub(r'Este código crea.*?visualizar el diagrama\.', '', texto_ia, flags=re.DOTALL | re.IGNORECASE)
    
    # Eliminar explicaciones comunes sobre el código HTML que la IA añade al final
    texto_limpio = re.sub(r'El código HTML proporcionado genera.*?aún más:', '', texto_limpio, flags=re.DOTALL | re.IGNORECASE)
    
    # Eliminar cualquier bloque de código JSON que pueda haberse colado
    texto_limpio = re.sub(r'```json\s*\{.*?\}\s*```', '', texto_limpio, flags=re.DOTALL)

    # Eliminar los marcadores de bloque de código de texto plano o html
    texto_limpio = re.sub(r'```(html|mermaid)?', '', texto_limpio)
    
    # Eliminar frases introductorias o de cierre que a veces añade la IA
    frases_a_eliminar = [
        r'^\s*Aquí tienes el contenido.*?:',
        r'^\s*Claro, aquí está la redacción para.*?:',
        r'^\s*##\s*.*?$' # Elimina cualquier título Markdown que la IA pueda repetir
    ]
    for patron in frases_a_eliminar:
        texto_limpio = re.sub(patron, '', texto_limpio, flags=re.IGNORECASE | re.MULTILINE)

    return texto_limpio.strip()
# =============================================================================
#           VERSIÓN DEFINITIVA DE phase_1_page()
# =============================================================================
def phase_1_page(model):
    """Página de Fase 1 que lee/escribe en subcarpetas y gestiona el estado correctamente."""
    if not st.session_state.get('selected_project'):
        st.warning("No se ha seleccionado ningún proyecto. Volviendo a la selección.")
        go_to_project_selection()
        st.rerun()

    project_name = st.session_state.selected_project['name']
    project_folder_id = st.session_state.selected_project['id']
    service = st.session_state.drive_service

    st.markdown(f"<h3>FASE 1: Análisis y Estructura</h3>", unsafe_allow_html=True)
    st.info(f"Estás trabajando en el proyecto: **{project_name}**")

    pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
    document_files = get_files_in_project(service, pliegos_folder_id)
    
    if document_files:
        st.success("Hemos encontrado estos archivos en la carpeta 'Pliegos' de tu proyecto:")
        with st.container(border=True):
            for file in document_files:
                cols = st.columns([4, 1])
                cols[0].write(f"📄 **{file['name']}**")
                if cols[1].button("Eliminar", key=f"del_{file['id']}", type="secondary"):
                    with st.spinner(f"Eliminando '{file['name']}'..."):
                        if delete_file_from_drive(service, file['id']):
                            st.toast(f"Archivo '{file['name']}' eliminado.")
                            st.rerun()
    else:
        st.info("La carpeta 'Pliegos' de este proyecto está vacía. Sube los archivos base.")

    with st.expander("Añadir o reemplazar documentación en la carpeta 'Pliegos'", expanded=not document_files):
        with st.container(border=True):
            st.subheader("Subir nuevos documentos")
            new_files_uploader = st.file_uploader("Arrastra aquí los nuevos Pliegos o Plantilla", type=['docx', 'pdf'], accept_multiple_files=True, key="new_files_uploader")
            if st.button("Guardar nuevos archivos en Drive"):
                if new_files_uploader:
                    with st.spinner("Subiendo archivos a la carpeta 'Pliegos'..."):
                        for file_obj in new_files_uploader:
                            upload_file_to_drive(service, file_obj, pliegos_folder_id)
                        st.rerun()
                else:
                    st.warning("Por favor, selecciona al menos un archivo para subir.")

    st.markdown("---")
    st.header("Análisis y Generación de Índice")
    
    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
    saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Cargar último índice generado", use_container_width=True, disabled=not saved_index_id):
            with st.spinner("Cargando índice desde Drive..."):
                index_content_bytes = download_file_from_drive(service, saved_index_id)
                index_data = json.loads(index_content_bytes.getvalue().decode('utf-8'))
                st.session_state.generated_structure = index_data
                st.session_state.uploaded_pliegos = document_files
                go_to_phase1_results()
                st.rerun()

    with col2:
        # Simplificamos este botón para que llame a la nueva función
        if st.button("Analizar Archivos y Generar Nuevo Índice", type="primary", use_container_width=True, disabled=not document_files):
            if handle_full_regeneration(model):
                go_to_phase1_results()
                st.rerun()

    st.write("")
    st.markdown("---")
    st.button("← Volver a Selección de Proyecto", on_click=back_to_project_selection_and_cleanup, use_container_width=True, key="back_to_projects")

# =============================================================================
#           VERSIÓN FINAL Y COMPLETA DE phase_1_results_page() CON SINCRONIZACIÓN
# =============================================================================

def phase_1_results_page(model):
    """Página para revisar, regenerar, ACEPTAR el índice y SINCRONIZAR carpetas."""
    st.markdown("<h3>FASE 1: Revisión de Resultados</h3>", unsafe_allow_html=True)
    st.markdown("Revisa el índice. Puedes hacer ajustes con feedback, regenerarlo todo desde cero, o aceptarlo para continuar.")
    st.markdown("---")
    st.button("← Volver a la gestión de archivos", on_click=go_to_phase1)

    if 'generated_structure' not in st.session_state or not st.session_state.generated_structure:
        st.warning("No se ha generado ninguna estructura.")
        return

    def handle_regeneration_with_feedback():
        feedback_text = st.session_state.feedback_area
        if not feedback_text:
            st.warning("Por favor, escribe tus indicaciones en el área de texto.")
            return

        with st.spinner("🧠 Incorporando tu feedback y regenerando la estructura..."):
            try:
                contenido_ia_regeneracion = [PROMPT_REGENERACION]
                contenido_ia_regeneracion.append("--- INSTRUCCIONES DEL USUARIO ---\n" + feedback_text)
                contenido_ia_regeneracion.append("--- ESTRUCTURA JSON ANTERIOR A CORREGIR ---\n" + json.dumps(st.session_state.generated_structure, indent=2))
                
                if st.session_state.get('uploaded_pliegos'):
                    service = st.session_state.drive_service
                    for file_info in st.session_state.uploaded_pliegos:
                        file_content_bytes = download_file_from_drive(service, file_info['id'])
                        contenido_ia_regeneracion.append({"mime_type": file_info['mimeType'], "data": file_content_bytes.getvalue()})

                generation_config = genai.GenerationConfig(response_mime_type="application/json")
                response_regeneracion = model.generate_content(contenido_ia_regeneracion, generation_config=generation_config)
                json_limpio_str_regenerado = limpiar_respuesta_json(response_regeneracion.text)
                
                if json_limpio_str_regenerado:
                    st.session_state.generated_structure = json.loads(json_limpio_str_regenerado)
                    st.toast("¡Estructura regenerada con feedback!")
                    st.session_state.feedback_area = "" # Limpiamos el área de texto
                else:
                    st.error("La IA no devolvió una estructura válida tras la regeneración.")
            except Exception as e:
                st.error(f"Ocurrió un error durante la regeneración: {e}")

    with st.container(border=True):
        mostrar_indice_desplegable(st.session_state.generated_structure.get('estructura_memoria'))
        st.markdown("---")
        st.subheader("Validación y Siguiente Paso")
        
        st.text_area("Si necesitas cambios, indícalos aquí:", key="feedback_area", placeholder="Ej: 'Une los apartados 1.1 y 1.2 en uno solo.'")
        
        col1, col2 = st.columns(2)
        with col1:
            st.button(
                "Regenerar con Feedback", 
                on_click=handle_regeneration_with_feedback, 
                use_container_width=True, 
                disabled=not st.session_state.get("feedback_area")
            )
        with col2:
            st.button(
                "🔁 Regenerar Índice Entero", 
                on_click=handle_full_regeneration, 
                args=(model,), 
                use_container_width=True, 
                help="Descarta este índice y genera uno nuevo desde cero analizando los pliegos otra vez."
            )

        if st.button("Aceptar Índice y Pasar a Fase 2 →", type="primary", use_container_width=True):
            with st.spinner("Sincronizando carpetas y guardando índice final en Drive..."):
                try:
                    service = st.session_state.drive_service
                    project_folder_id = st.session_state.selected_project['id']
                    
                    deleted_count = sync_guiones_folders_with_index(service, project_folder_id, st.session_state.generated_structure)
                    if deleted_count > 0:
                        st.success(f"Limpieza completada: {deleted_count} carpetas de guiones obsoletas eliminadas.")

                    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
                    json_bytes = json.dumps(st.session_state.generated_structure, indent=2).encode('utf-8')
                    mock_file_obj = io.BytesIO(json_bytes)
                    mock_file_obj.name = "ultimo_indice.json"
                    mock_file_obj.type = "application/json"
                    
                    saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
                    if saved_index_id:
                        delete_file_from_drive(service, saved_index_id)
                    upload_file_to_drive(service, mock_file_obj, docs_app_folder_id)
                    st.toast("Índice final guardado en tu proyecto de Drive.")
                    
                    go_to_phase2()
                    st.rerun()

                except Exception as e:
                    st.error(f"Ocurrió un error durante la sincronización o guardado: {e}")

# =============================================================================
#           VERSIÓN FINAL Y OPTIMIZADA DE phase_2_page (CORRIGE TIMEOUT)
# =============================================================================

def phase_2_page(model):
    """Centro de mando para la generación de guiones con opciones individuales y en lote."""
    st.markdown("<h3>FASE 2: Centro de Mando de Guiones</h3>", unsafe_allow_html=True)
    st.markdown("Gestiona tus guiones de forma individual o selecciónalos para generarlos en lote.")
    st.markdown("---")

    # --- SETUP INICIAL Y CARGA DE ÍNDICE ---
    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    if 'generated_structure' not in st.session_state:
        st.info("Sincronizando índice desde Google Drive...")
        try:
            docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
            saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
            if saved_index_id:
                index_content_bytes = download_file_from_drive(service, saved_index_id)
                st.session_state.generated_structure = json.loads(index_content_bytes.getvalue().decode('utf-8'))
                st.rerun()
            else:
                st.warning("No se ha encontrado un índice guardado. Por favor, vuelve a la Fase 1 para generar uno.")
                if st.button("← Ir a Fase 1"): go_to_phase1(); st.rerun()
                return
        except Exception as e:
            st.error(f"Error al cargar el índice desde Drive: {e}")
            return

    # --- CONSTRUCCIÓN DE LISTA ROBUSTA ---
    estructura = st.session_state.generated_structure.get('estructura_memoria', [])
    matices_originales = st.session_state.generated_structure.get('matices_desarrollo', [])
    matices_dict = {item.get('subapartado', ''): item for item in matices_originales if isinstance(item, dict) and 'subapartado' in item}
    if not estructura: st.error("La estructura JSON no contiene la clave 'estructura_memoria'."); return
    subapartados_a_mostrar = []
    for seccion in estructura:
        apartado_principal = seccion.get('apartado', 'Sin Título')
        for subapartado_titulo in seccion.get('subapartados', []):
            matiz_existente = matices_dict.get(subapartado_titulo)
            if matiz_existente: subapartados_a_mostrar.append(matiz_existente)
            else: subapartados_a_mostrar.append({"apartado": apartado_principal, "subapartado": subapartado_titulo, "indicaciones": "No se encontraron indicaciones detalladas."})
    if not subapartados_a_mostrar: st.warning("El índice no contiene subapartados."); return

    # --- FUNCIONES DE ACCIÓN INTERNAS ---
    def ejecutar_generacion(titulo, indicaciones_completas, show_toast=True):
        nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", titulo)
        nombre_archivo = nombre_limpio + ".docx"
        try:
            guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
            subapartado_guion_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_folder_id)
            pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
            pliegos_en_drive = get_files_in_project(service, pliegos_folder_id)
            contenido_ia = [PROMPT_PREGUNTAS_TECNICAS_INDIVIDUAL]
            contenido_ia.append("--- INDICACIONES PARA ESTE APARTADO ---\n" + json.dumps(indicaciones_completas, indent=2, ensure_ascii=False))
            for file_info in pliegos_en_drive:
                file_content_bytes = download_file_from_drive(service, file_info['id'])
                contenido_ia.append({"mime_type": file_info['mimeType'], "data": file_content_bytes.getvalue()})
            doc_extra_key = f"upload_{titulo}"
            if doc_extra_key in st.session_state and st.session_state[doc_extra_key]:
                for uploaded_file in st.session_state[doc_extra_key]:
                    contenido_ia.append("--- DOCUMENTACIÓN DE APOYO ADICIONAL ---\n")
                    contenido_ia.append({"mime_type": uploaded_file.type, "data": uploaded_file.getvalue()})
                    upload_file_to_drive(service, uploaded_file, subapartado_guion_folder_id)
            response = model.generate_content(contenido_ia)
            documento = docx.Document()
            agregar_markdown_a_word(documento, response.text)
            doc_io = io.BytesIO()
            documento.save(doc_io)
            word_file_obj = io.BytesIO(doc_io.getvalue())
            word_file_obj.name = nombre_archivo
            word_file_obj.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            upload_file_to_drive(service, word_file_obj, subapartado_guion_folder_id)
            if show_toast: st.toast(f"Borrador para '{titulo}' generado y guardado.")
            return True
        except Exception as e: st.error(f"Error al generar '{titulo}': {e}"); return False

    def ejecutar_regeneracion(titulo, file_id_borrador): pass # Tu código original aquí
    def ejecutar_borrado(titulo, folder_id_to_delete): pass # Tu código original aquí

    # =============================================================================
    #           OPTIMIZACIÓN CLAVE: OBTENER DATOS DE DRIVE UNA SOLA VEZ
    # =============================================================================
    with st.spinner("Sincronizando con Google Drive..."):
        guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
        # Hacemos UNA sola llamada a la API para obtener todas las carpetas existentes
        carpetas_existentes_response = get_files_in_project(service, guiones_folder_id)
        carpetas_de_guiones_existentes = {f['name']: f['id'] for f in carpetas_existentes_response if f['mimeType'] == 'application/vnd.google-apps.folder'}
        nombres_carpetas_existentes = set(carpetas_de_guiones_existentes.keys())

    # =============================================================================
    #           SECCIÓN SUPERIOR PARA ACCIONES EN LOTE (AHORA OPTIMIZADA)
    # =============================================================================
    st.subheader("Generación de Borradores en Lote")
    
    pending_keys = [matiz.get('subapartado') for matiz in subapartados_a_mostrar if re.sub(r'[\\/*?:"<>|]', "", matiz.get('subapartado')) not in nombres_carpetas_existentes]
    
    def toggle_all_checkboxes():
        new_state = st.session_state.select_all_checkbox
        for key in pending_keys: st.session_state[f"cb_{key}"] = new_state

    with st.container(border=True):
        col_sel_1, col_sel_2 = st.columns([1, 2])
        with col_sel_1:
            st.checkbox("Seleccionar Todos / Ninguno", key="select_all_checkbox", on_change=toggle_all_checkboxes, disabled=not pending_keys)
        with col_sel_2:
            selected_keys = [key for key in pending_keys if st.session_state.get(f"cb_{key}")]
            num_selected = len(selected_keys)
            if st.button(f"🚀 Generar {num_selected} borradores seleccionados", type="primary", use_container_width=True, disabled=(num_selected == 0)):
                progress_bar = st.progress(0, text="Iniciando generación en lote...")
                items_to_generate = [matiz for matiz in subapartados_a_mostrar if matiz.get('subapartado') in selected_keys]
                for i, matiz_a_generar in enumerate(items_to_generate):
                    titulo = matiz_a_generar.get('subapartado')
                    progress_text = f"Generando ({i+1}/{num_selected}): {titulo}"
                    progress_bar.progress((i + 1) / num_selected, text=progress_text)
                    ejecutar_generacion(titulo, matiz_a_generar, show_toast=False)
                progress_bar.progress(1.0, text="¡Generación en lote completada!")
                st.success(f"{num_selected} borradores generados.")
                st.balloons()
                st.rerun()

    st.markdown("---")
    st.subheader("Gestión de Guiones de Subapartados")

    # =============================================================================
    #           INTERFAZ DE GESTIÓN DE GUIONES (DISEÑO HÍBRIDO Y OPTIMIZADO)
    # =============================================================================
    for i, matiz in enumerate(subapartados_a_mostrar):
        subapartado_titulo = matiz.get('subapartado')
        if not subapartado_titulo: continue
        nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", subapartado_titulo)
        
        if nombre_limpio in nombres_carpetas_existentes:
            estado = "📄 Generado"
            subapartado_folder_id = carpetas_de_guiones_existentes[nombre_limpio]
            files_in_subfolder = get_files_in_project(service, subapartado_folder_id)
            file_info = next((f for f in files_in_subfolder if f['name'].endswith('.docx')), None)
        else:
            estado = "⚪ No Generado"
            file_info, subapartado_folder_id = None, None

        with st.container(border=True):
            col1, col2 = st.columns([2, 1])
            with col1:
                if estado == "⚪ No Generado": st.checkbox(f"**{subapartado_titulo}**", key=f"cb_{subapartado_titulo}")
                else: st.write(f"**{subapartado_titulo}**")
                st.caption(f"Estado: {estado}")
                if estado == "⚪ No Generado":
                    st.file_uploader("Aportar documentación de apoyo", type=['pdf', 'docx', 'txt'], key=f"upload_{subapartado_titulo}", accept_multiple_files=True, label_visibility="collapsed")
            with col2:
                if estado == "📄 Generado" and file_info:
                    link = f"https://docs.google.com/document/d/{file_info['id']}/edit"
                    st.link_button("Revisar en Drive", link, use_container_width=True)
                    if st.button("Re-Generar con Feedback", key=f"regen_{i}", type="primary", use_container_width=True):
                        ejecutar_regeneracion(subapartado_titulo, file_info['id'])
                    if st.button("🗑️ Borrar", key=f"del_{i}", use_container_width=True):
                         ejecutar_borrado(subapartado_titulo, subapartado_folder_id)
                else:
                    if st.button("Generar Borrador", key=f"gen_{i}", use_container_width=True):
                        with st.spinner(f"Generando borrador para '{subapartado_titulo}'..."):
                            if ejecutar_generacion(subapartado_titulo, matiz): st.rerun()

    # --- NAVEGACIÓN ---
    st.markdown("---")
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Revisión de Índice (F1)", on_click=go_to_phase1_results, use_container_width=True)
    with col_nav2:
        st.button("Ir a Plan de Prompts (F3) →", on_click=go_to_phase3, use_container_width=True)

# =============================================================================
#           VERSIÓN FINAL Y OPTIMIZADA DE phase_3_page (CON SELECCIÓN MÚLTIPLE)
# =============================================================================

# =============================================================================
#           VERSIÓN FINAL Y OPTIMIZADA DE phase_3_page (CON BOTÓN DE BORRADO)
# =============================================================================

# =============================================================================
#           VERSIÓN FINAL Y COMPLETA DE phase_3_page (CORRECCIÓN DE UNIFICACIÓN)
# =============================================================================

def phase_3_page(model):
    """Página interactiva para generar, borrar, descargar y unificar planes de prompts."""
    st.markdown("<h3>FASE 3: Centro de Mando de Prompts</h3>", unsafe_allow_html=True)
    st.markdown("Genera planes de prompts de forma individual o selecciónalos para procesarlos en lote.")
    st.markdown("---")

    # --- SETUP INICIAL Y CARGA DE ÍNDICE ---
    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)

    if 'generated_structure' not in st.session_state:
        st.info("Sincronizando índice desde Google Drive...")
        saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
        if saved_index_id:
            index_content_bytes = download_file_from_drive(service, saved_index_id)
            st.session_state.generated_structure = json.loads(index_content_bytes.getvalue().decode('utf-8'))
            st.rerun()
        else:
            st.warning("No se ha encontrado un índice. Vuelve a Fase 1 para generarlo.")
            if st.button("← Ir a Fase 1"): go_to_phase1(); st.rerun()
            return

    # --- CONSTRUCCIÓN DE LISTA ROBUSTA ---
    estructura = st.session_state.generated_structure.get('estructura_memoria', [])
    matices_originales = st.session_state.generated_structure.get('matices_desarrollo', [])
    matices_dict = {item.get('subapartado', ''): item for item in matices_originales if isinstance(item, dict) and 'subapartado' in item}
    if not estructura: st.error("La estructura JSON no contiene la clave 'estructura_memoria'."); return
    subapartados_a_mostrar = []
    for seccion in estructura:
        apartado_principal = seccion.get('apartado', 'Sin Título')
        for subapartado_titulo in seccion.get('subapartados', []):
            matiz_existente = matices_dict.get(subapartado_titulo)
            if matiz_existente: subapartados_a_mostrar.append(matiz_existente)
            else: subapartados_a_mostrar.append({"apartado": apartado_principal, "subapartado": subapartado_titulo, "indicaciones": "No se encontraron indicaciones detalladas."})
    if not subapartados_a_mostrar: st.warning("El índice no contiene subapartados."); return

    # --- FUNCIONES DE ACCIÓN INTERNAS ---
    def handle_individual_generation(matiz_info, callback_model, show_toast=True):
        apartado_titulo = matiz_info.get("apartado", "N/A"); subapartado_titulo = matiz_info.get("subapartado", "N/A")
        # ... (código interno de esta función sin cambios)
        json_limpio_str = ""
        try:
            guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
            nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", subapartado_titulo)
            subapartado_folder_id = find_or_create_folder(service, nombre_limpio, parent_id=guiones_main_folder_id)
            contexto_adicional_str = ""; files_in_subfolder = get_files_in_project(service, subapartado_folder_id)
            for file_info in files_in_subfolder:
                file_bytes = download_file_from_drive(service, file_info['id'])
                if file_info['name'].endswith('.docx'):
                    doc = docx.Document(io.BytesIO(file_bytes.getvalue())); texto_doc = "\n".join([p.text for p in doc.paragraphs])
                    contexto_adicional_str += f"\n--- CONTENIDO DEL GUION ({file_info['name']}) ---\n{texto_doc}\n"
                elif file_info['name'].endswith('.pdf'):
                    reader = PdfReader(io.BytesIO(file_bytes.getvalue())); texto_pdf = "".join(page.extract_text() for page in reader.pages)
                    contexto_adicional_str += f"\n--- CONTENIDO DEL PDF DE APOYO ({file_info['name']}) ---\n{texto_pdf}\n"
            pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
            pliegos_files_info = get_files_in_project(service, pliegos_folder_id)
            pliegos_content_for_ia = [{"mime_type": f['mimeType'], "data": download_file_from_drive(service, f['id']).getvalue()} for f in pliegos_files_info]
            prompt_final = PROMPT_DESARROLLO.format(apartado_titulo=apartado_titulo, subapartado_titulo=subapartado_titulo, indicaciones=matiz_info.get("indicaciones", ""))
            contenido_ia = [prompt_final] + pliegos_content_for_ia
            if contexto_adicional_str: contenido_ia.append("--- CONTEXTO ADICIONAL DE GUIONES Y DOCUMENTACIÓN DE APOYO ---\n" + contexto_adicional_str)
            generation_config = genai.GenerationConfig(response_mime_type="application/json")
            response = callback_model.generate_content(contenido_ia, generation_config=generation_config)
            json_limpio_str = limpiar_respuesta_json(response.text)
            if json_limpio_str:
                json_sanitizado = sanitize_json_string(json_limpio_str)
                plan_parcial_obj = json.loads(json_sanitizado)
                json_bytes = json.dumps(plan_parcial_obj, indent=2, ensure_ascii=False).encode('utf-8')
                mock_file_obj = io.BytesIO(json_bytes); mock_file_obj.name = "prompts_individual.json"; mock_file_obj.type = "application/json"
                old_plan_id = find_file_by_name(service, "prompts_individual.json", subapartado_folder_id)
                if old_plan_id: delete_file_from_drive(service, old_plan_id)
                upload_file_to_drive(service, mock_file_obj, subapartado_folder_id)
                if show_toast: st.toast(f"Plan para '{subapartado_titulo}' guardado.")
                st.rerun()
                return True
        except json.JSONDecodeError as json_err:
             st.error(f"Error Crítico: La IA devolvió un JSON inválido para '{subapartado_titulo}' que no se pudo reparar. Detalles: {json_err}")
             st.code(json_limpio_str)
             return False
        except Exception as e:
            st.error(f"Error generando prompts para '{subapartado_titulo}': {e}")
            return False

    def handle_individual_deletion(titulo, plan_id_to_delete):
        """Elimina un archivo de plan individual y refresca la página."""
        with st.spinner(f"Eliminando el plan para '{titulo}'..."):
            if delete_file_from_drive(service, plan_id_to_delete):
                st.toast(f"Plan para '{titulo}' eliminado con éxito.")
                st.rerun()

    # =============== ¡INICIO DE LA CORRECCIÓN! ===============
    def handle_conjunto_generation():
        """Unifica todos los planes individuales en un único archivo maestro."""
        with st.spinner("Unificando todos los planes de prompts..."):
            try:
                # 1. Localizar la carpeta principal de guiones
                guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
                
                # 2. Listar todas las subcarpetas de guiones
                carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
                
                plan_conjunto_final = {"plan_de_prompts": []}
                
                # 3. Iterar sobre cada subcarpeta para encontrar y leer su plan individual
                for nombre_carpeta, folder_id in carpetas_de_guiones.items():
                    plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
                    if plan_id:
                        # 4. Descargar y procesar el JSON individual
                        json_bytes = download_file_from_drive(service, plan_id).getvalue()
                        plan_individual_obj = json.loads(json_bytes.decode('utf-8'))
                        
                        # 5. Añadir los prompts del plan individual a la lista maestra
                        prompts_de_este_plan = plan_individual_obj.get("plan_de_prompts", [])
                        plan_conjunto_final["plan_de_prompts"].extend(prompts_de_este_plan)

                if not plan_conjunto_final["plan_de_prompts"]:
                    st.warning("No se encontraron planes individuales para unificar. Genera al menos uno.")
                    return

                # 6. Preparar el archivo unificado para subirlo a Drive
                nombre_archivo_final = "plan_de_prompts_conjunto.json"
                json_bytes_finales = json.dumps(plan_conjunto_final, indent=2, ensure_ascii=False).encode('utf-8')
                
                mock_file_obj = io.BytesIO(json_bytes_finales)
                mock_file_obj.name = nombre_archivo_final
                mock_file_obj.type = "application/json"
                
                # 7. Borrar el archivo antiguo si existe y subir el nuevo
                old_conjunto_id = find_file_by_name(service, nombre_archivo_final, docs_app_folder_id)
                if old_conjunto_id:
                    delete_file_from_drive(service, old_conjunto_id)
                
                upload_file_to_drive(service, mock_file_obj, docs_app_folder_id)
                
                st.success(f"¡Plan conjunto generado y guardado! Se unificaron {len(plan_conjunto_final['plan_de_prompts'])} prompts.")
                st.balloons()

            except Exception as e:
                st.error(f"Ocurrió un error durante la unificación: {e}")
    # =============== ¡FIN DE LA CORRECCIÓN! ===============

    # OPTIMIZACIÓN: OBTENER ESTADO DE PLANES UNA SOLA VEZ
    with st.spinner("Verificando estado de los planes de prompts..."):
        # ... (código sin cambios)
        guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
        carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
        planes_individuales_existentes = {}
        for nombre_carpeta, folder_id in carpetas_de_guiones.items():
            plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
            if plan_id: planes_individuales_existentes[nombre_carpeta] = plan_id

    # SECCIÓN SUPERIOR PARA ACCIONES EN LOTE
    st.subheader("Generación de Planes de Prompts en Lote")
    # ... (código sin cambios)
    pending_keys = [matiz.get('subapartado') for matiz in subapartados_a_mostrar if re.sub(r'[\\/*?:"<>|]', "", matiz.get('subapartado')) in carpetas_de_guiones and re.sub(r'[\\/*?:"<>|]', "", matiz.get('subapartado')) not in planes_individuales_existentes]
    def toggle_all_prompt_checkboxes():
        new_state = st.session_state.select_all_prompts_checkbox
        for key in pending_keys: st.session_state[f"pcb_{key}"] = new_state
    with st.container(border=True):
        col_sel_1, col_sel_2 = st.columns([1, 2])
        with col_sel_1:
            st.checkbox("Seleccionar Todos / Ninguno", key="select_all_prompts_checkbox", on_change=toggle_all_prompt_checkboxes, disabled=not pending_keys)
        with col_sel_2:
            selected_keys = [key for key in pending_keys if st.session_state.get(f"pcb_{key}")]
            num_selected = len(selected_keys)
            if st.button(f"🚀 Generar {num_selected} planes seleccionados", type="primary", use_container_width=True, disabled=(num_selected == 0)):
                progress_bar = st.progress(0, text="Iniciando generación en lote de planes...")
                items_to_generate = [matiz for matiz in subapartados_a_mostrar if matiz.get('subapartado') in selected_keys]
                for i, matiz_a_generar in enumerate(items_to_generate):
                    titulo = matiz_a_generar.get('subapartado')
                    progress_text = f"Generando plan ({i+1}/{num_selected}): {titulo}"
                    progress_bar.progress((i + 1) / num_selected, text=progress_text)
                    handle_individual_generation(matiz_a_generar, model, show_toast=False)
                progress_bar.progress(1.0, text="¡Generación en lote completada!")
                st.success(f"{num_selected} planes de prompts generados.")
                st.balloons()
                st.rerun()

    st.markdown("---")
    st.subheader("Gestión de Planes de Prompts")

    # INTERFAZ DE GESTIÓN (HÍBRIDA Y OPTIMIZADA)
    # ... (código sin cambios)
    for i, matiz in enumerate(subapartados_a_mostrar):
        subapartado_titulo = matiz.get("subapartado");
        if not subapartado_titulo: continue
        nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", subapartado_titulo)
        guion_generado = nombre_limpio in carpetas_de_guiones
        plan_individual_id = planes_individuales_existentes.get(nombre_limpio)
        with st.container(border=True):
            col1, col2 = st.columns([2, 1])
            with col1:
                if not plan_individual_id and guion_generado:
                    st.checkbox(f"**{subapartado_titulo}**", key=f"pcb_{subapartado_titulo}")
                else: st.write(f"**{subapartado_titulo}**")
                if not guion_generado: st.warning("⚠️ Guion no generado en Fase 2. No se puede crear un plan.")
                elif plan_individual_id:
                    st.success("✔️ Plan generado")
                    with st.expander("Ver / Descargar Plan Individual"):
                        json_bytes = download_file_from_drive(service, plan_individual_id).getvalue()
                        st.json(json_bytes.decode('utf-8'))
                        st.download_button("Descargar JSON", data=json_bytes, file_name=f"prompts_{nombre_limpio}.json", mime="application/json", key=f"dl_{i}")
                else: st.info("⚪ Pendiente de generar plan de prompts")
            with col2:
                if not plan_individual_id:
                    st.button("Generar Plan de Prompts", key=f"gen_ind_{i}", on_click=handle_individual_generation, args=(matiz, model, True), use_container_width=True, type="primary", disabled=not guion_generado)
                else:
                    st.button("Re-generar Plan", key=f"gen_regen_{i}", on_click=handle_individual_generation, args=(matiz, model, True), use_container_width=True, type="secondary")
                    st.button("🗑️ Borrar Plan", key=f"del_plan_{i}", on_click=handle_individual_deletion, args=(subapartado_titulo, plan_individual_id), use_container_width=True)

    # BOTONES DE NAVEGACIÓN Y ACCIÓN FINAL
    st.markdown("---")
    st.button("🚀 Unificar y Guardar Plan de Prompts Conjunto", on_click=handle_conjunto_generation, use_container_width=True, type="primary", help="Unifica todos los planes individuales generados en un único archivo maestro.")
    col_nav3_1, col_nav3_2 = st.columns(2)
    with col_nav3_1:
        st.button("← Volver al Centro de Mando (F2)", on_click=go_to_phase2, use_container_width=True)
    with col_nav3_2:
        st.button("Ir a Redacción Final (F4) →", on_click=go_to_phase4, use_container_width=True)
# =============================================================================
#           FASE 4 - REDACCIÓN Y ENSAMBLAJE FINAL
# =============================================================================

import re
import io
import os
import time
import docx
import json
import streamlit as st
from pypdf import PdfReader # Asegúrate de que esta importación esté al principio de tu script
import imgkit # Y esta también

# =============================================================================
#           FASE 4 - REDACCIÓN Y ENSAMBLAJE FINAL (VERSIÓN CORREGIDA)
# =============================================================================
# =============================================================================
#           FASE 4 - VERSIÓN ACTUALIZADA PARA PREPARAR BORRADOR
# =============================================================================

def phase_4_page(model):
    """Página para ejecutar el plan de prompts y generar el borrador inicial del documento Word."""
    st.markdown("<h3>FASE 4: Redacción del Borrador Inicial</h3>", unsafe_allow_html=True)
    st.markdown("Ejecuta el plan de prompts para generar el contenido completo de la memoria técnica. Este borrador se usará como base para el refinamiento final en la siguiente fase.")
    st.markdown("---")

    service = st.session_state.drive_service
    project_folder_id = st.session_state.selected_project['id']
    docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
    plan_conjunto_id = find_file_by_name(service, "plan_de_prompts_conjunto.json", docs_app_folder_id)

    if not plan_conjunto_id:
        st.warning("No se ha encontrado un 'plan_de_prompts_conjunto.json'. Vuelve a la Fase 3 para generarlo.")
        if st.button("← Ir a Fase 3"): go_to_phase3(); st.rerun()
        return

    try:
        json_bytes = download_file_from_drive(service, plan_conjunto_id).getvalue()
        plan_de_accion = json.loads(json_bytes.decode('utf-8'))
        lista_de_prompts = plan_de_accion.get("plan_de_prompts", [])
        lista_de_prompts.sort(key=lambda x: x.get('prompt_id', ''))
        st.success(f"✔️ Plan de acción cargado. Se ejecutarán {len(lista_de_prompts)} prompts para crear el borrador.")
    except Exception as e:
        st.error(f"Error al cargar o procesar el plan de acción: {e}"); return

    if 'generated_doc_buffer' not in st.session_state: st.session_state.generated_doc_buffer = None
    if 'generated_doc_filename' not in st.session_state: st.session_state.generated_doc_filename = ""

    button_text = "🔁 Volver a Generar Borrador" if st.session_state.generated_doc_buffer else "🚀 Iniciar Redacción y Generar Borrador"
    if st.button(button_text, type="primary", use_container_width=True):
        if not lista_de_prompts:
            st.warning("El plan de acción está vacío."); return

        generation_successful = False
        documento = docx.Document()

        try:
            with st.spinner("Iniciando redacción del borrador... Esto puede tardar varios minutos."):
                chat_redaccion = model.start_chat()
                progress_bar = st.progress(0, text="Configurando sesión de chat...")
                
                ultimo_apartado_escrito = None
                ultimo_subapartado_escrito = None
                
                for i, tarea in enumerate(lista_de_prompts):
                    progress_text = f"Procesando Tarea {i+1}/{len(lista_de_prompts)}: {tarea.get('subapartado_referencia', 'N/A')}"
                    progress_bar.progress((i + 1) / len(lista_de_prompts), text=progress_text)
                    
                    apartado_actual = tarea.get("apartado_referencia")
                    subapartado_actual = tarea.get("subapartado_referencia")
                    if apartado_actual and apartado_actual != ultimo_apartado_escrito:
                        if ultimo_apartado_escrito is not None: documento.add_page_break()
                        documento.add_heading(apartado_actual, level=1)
                        ultimo_apartado_escrito = apartado_actual
                        ultimo_subapartado_escrito = None
                    if subapartado_actual and subapartado_actual != ultimo_subapartado_escrito:
                        documento.add_heading(subapartado_actual, level=2)
                        ultimo_subapartado_escrito = subapartado_actual
                    
                    respuesta_ia_bruta = ""
                    prompt_actual = tarea.get("prompt_para_asistente")
                    if prompt_actual:
                        response = chat_redaccion.send_message(prompt_actual)
                        respuesta_ia_bruta = response.text
                        time.sleep(1)
                    
                    if respuesta_ia_bruta:
                        patron_html = re.compile(r'```html\s*([\s\S]*?)\s*```|(<div[\s\S]*<\/div>)|(<!DOCTYPE html>[\s\S]*?<\/html>)', re.DOTALL)
                        match_html = patron_html.search(respuesta_ia_bruta)

                        if match_html:
                            html_puro = next(g for g in match_html.groups() if g is not None)
                            texto_narrativo_completo = respuesta_ia_bruta[:match_html.start()] + respuesta_ia_bruta[match_html.end():]
                            
                            # APLICAMOS LAS NUEVAS FUNCIONES DE LIMPIEZA
                            texto_limpio = limpiar_respuesta_final(texto_narrativo_completo)
                            texto_corregido = corregir_numeracion_markdown(texto_limpio)
                            
                            if texto_corregido:
                                agregar_markdown_a_word(documento, texto_corregido)
                            
                            image_file = html_a_imagen(wrap_html_fragment(html_puro), f"temp_img_{i}.png")
                            if image_file and os.path.exists(image_file):
                                documento.add_picture(image_file, width=docx.shared.Inches(6.5))
                                os.remove(image_file)
                            else:
                                documento.add_paragraph("[ERROR AL GENERAR IMAGEN DESDE HTML]")
                        else:
                            # APLICAMOS LAS NUEVAS FUNCIONES DE LIMPIEZA
                            texto_limpio = limpiar_respuesta_final(respuesta_ia_bruta)
                            texto_corregido = corregir_numeracion_markdown(texto_limpio)
                            if texto_corregido:
                                agregar_markdown_a_word(documento, texto_corregido)
                
                generation_successful = True

        except Exception as e:
            st.error(f"Ocurrió un error crítico durante la generación del borrador: {e}")
        
        if generation_successful:
            project_name = st.session_state.selected_project['name']
            safe_project_name = re.sub(r'[\\/*?:"<>|]', "", project_name).replace(' ', '_')
            nombre_archivo_final = f"Memoria_Tecnica_{safe_project_name}_Borrador.docx"
            
            doc_io = io.BytesIO()
            documento.save(doc_io)
            doc_io.seek(0)
            
            st.session_state.generated_doc_buffer = doc_io
            st.session_state.generated_doc_filename = nombre_archivo_final
            
            st.success("¡Borrador inicial generado con éxito!")
            st.rerun()

    if st.session_state.generated_doc_buffer:
        st.info("El borrador inicial está listo. Ahora puedes descargarlo o pasar a la fase final de refinamiento.")
        st.download_button(
            label="📄 Descargar Borrador Inicial (.docx)",
            data=st.session_state.generated_doc_buffer,
            file_name=st.session_state.generated_doc_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown("---")
    # NAVEGACIÓN ACTUALIZADA
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Fase 3 (Plan de Prompts)", on_click=go_to_phase3, use_container_width=True)
    with col_nav2:
        # Botón clave para avanzar a la nueva fase. Se activa solo si hay un borrador.
        st.button("Ir a Refinamiento Final (F5) →", on_click=go_to_phase5, use_container_width=True, type="primary", disabled=not st.session_state.generated_doc_buffer)

# =============================================================================
#           FASE 5 - VERSIÓN FINAL CON INTRODUCCIÓN Y COHESIÓN
# =============================================================================

def phase_5_page(model):
    """
    Fase final que primero genera una introducción estratégica y luego refina
    la cohesión y el estilo del documento completo para crear la versión definitiva.
    """
    st.markdown("<h3>FASE 5: Refinamiento y Ensamblaje Definitivo</h3>", unsafe_allow_html=True)
    st.markdown("Este es el último paso. El asistente realizará dos tareas clave:")
    st.markdown("1.  **Creará una introducción** estratégica basada en el contenido completo.")
    st.markdown("2.  **Revisará todo el documento** para mejorar el flujo, añadir referencias entre apartados y garantizar una voz coherente.")
    st.info("Este proceso analiza todo el texto y realiza dos llamadas a la IA, por lo que puede tardar varios minutos.")
    st.markdown("---")

    # Comprobamos si tenemos un documento de la fase 4 para trabajar
    if 'generated_doc_buffer' not in st.session_state or not st.session_state.generated_doc_buffer:
        st.warning("No se ha encontrado un borrador de la Fase 4. Por favor, completa la fase anterior primero.")
        if st.button("← Ir a Fase 4"): go_to_phase4(); st.rerun()
        return

    if 'refined_doc_buffer' not in st.session_state: st.session_state.refined_doc_buffer = None

    if st.button("✨ Iniciar Ensamblaje y Refinamiento Final", type="primary", use_container_width=True):
        
        # Inicializamos variables para los resultados
        introduccion_markdown = ""
        cuerpo_refinado_markdown = ""
        proceso_exitoso = False

        try:
            # Extraer texto del documento Word generado en Fase 4
            buffer = st.session_state.generated_doc_buffer
            buffer.seek(0)
            documento_original = docx.Document(buffer)
            # Extraemos párrafos con texto para evitar líneas vacías excesivas
            texto_completo_original = "\n".join([p.text for p in documento_original.paragraphs if p.text.strip()])

            # --- PASO 1: GENERAR INTRODUCCIÓN ---
            with st.spinner("Paso 1/2: El estratega IA está redactando la introducción..."):
                response_intro = model.generate_content([PROMPT_GENERAR_INTRODUCCION, texto_completo_original])
                introduccion_markdown = response_intro.text
                st.toast("Introducción generada.")
                time.sleep(1) # Pequeña pausa

            # --- PASO 2: REFINAR EL CUERPO DEL DOCUMENTO ---
            with st.spinner("Paso 2/2: El editor IA está aplicando cohesión y referencias a todo el documento..."):
                response_cohesion = model.generate_content([PROMPT_COHESION_FINAL, texto_completo_original])
                cuerpo_refinado_markdown = response_cohesion.text
                st.toast("Cuerpo del documento refinado.")

            proceso_exitoso = True

        except Exception as e:
            st.error(f"Ocurrió un error crítico durante el refinamiento: {e}")

        if proceso_exitoso:
            # --- PASO 3: ENSAMBLAJE FINAL DEL DOCUMENTO WORD ---
            with st.spinner("Ensamblando la versión definitiva..."):
                documento_final = docx.Document()
                
                # Añadir la introducción
                documento_final.add_heading("Introducción", level=1)
                intro_limpia = limpiar_respuesta_final(introduccion_markdown)
                intro_corregida = corregir_numeracion_markdown(intro_limpia) # Aunque no suele haber listas, por si acaso
                agregar_markdown_a_word(documento_final, intro_corregida)
                
                documento_final.add_page_break()

                # Añadir el cuerpo refinado
                cuerpo_limpio = limpiar_respuesta_final(cuerpo_refinado_markdown)
                cuerpo_corregido = corregir_numeracion_markdown(cuerpo_limpio)
                agregar_markdown_a_word(documento_final, cuerpo_corregido)

                # Guardar en buffer para descarga y en Drive
                doc_io_final = io.BytesIO()
                documento_final.save(doc_io_final)
                doc_io_final.seek(0)

                st.session_state.refined_doc_buffer = doc_io_final
                original_filename = st.session_state.generated_doc_filename
                st.session_state.refined_doc_filename = original_filename.replace(".docx", "_Definitivo.docx")
                
                st.rerun()

    if st.session_state.refined_doc_buffer:
        st.balloons()
        st.success("¡Tu memoria técnica definitiva está lista!")
        st.download_button(
            label="🏆 Descargar Versión Definitiva (.docx)",
            data=st.session_state.refined_doc_buffer,
            file_name=st.session_state.refined_doc_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown("---")
    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        st.button("← Volver a Fase 4 (Borrador Inicial)", on_click=go_to_phase4, use_container_width=True)
    with col_nav2:
        st.button("↩️ Volver a Selección de Proyecto", on_click=back_to_project_selection_and_cleanup, use_container_width=True)
# =============================================================================
#                        LÓGICA PRINCIPAL (ROUTER)
# =============================================================================

credentials = get_credentials()

if not credentials:
    landing_page()
else:
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-1.5-pro-latest') 
    except Exception as e:
        st.error(f"Error al configurar la API de Gemini. Verifica tu 'GEMINI_API_KEY' en los secrets. Detalle: {e}")
        st.stop()
        
    # El router ahora incluye la nueva Fase 5
    if st.session_state.page == 'landing' or st.session_state.page == 'project_selection':
        project_selection_page()
    elif st.session_state.page == 'phase_1':
        phase_1_page(model)
    elif st.session_state.page == 'phase_1_results':
        phase_1_results_page(model)
    elif st.session_state.page == 'phase_2':
        phase_2_page(model)
    elif st.session_state.page == 'phase_3':
        phase_3_page(model)
    elif st.session_state.page == 'phase_4':
        phase_4_page(model)
    elif st.session_state.page == 'phase_5': # <-- NUEVA LÍNEA
        phase_5_page(model)               # <-- NUEVA LÍNEA
Y no olvides añadir las nuevas funciones de navegación y limpieza de estado que mencioné en la respuesta anterior si aún no lo has hecho. Te las incluyo aquí de nuevo por si acaso:
code
Python
# AÑADE ESTO A TU SECCIÓN DE NAVEGACIÓN Y GESTIÓN DE ESTADO

# ... (junto a tus otras funciones go_to_... )
def go_to_phase5(): st.session_state.page = 'phase_5'

# ... (en la misma sección, actualiza tu función de limpieza)
def back_to_project_selection_and_cleanup():
    # Asegúrate de que las nuevas variables de estado se limpien también
    for key in ['generated_structure', 'word_file', 'uploaded_template', 
                'uploaded_pliegos', 'selected_project', 'generated_doc_buffer', 
                'refined_doc_buffer', 'generated_doc_filename', 'refined_doc_filename']:
        if key in st.session_state: 
            del st.session_state[key]
    go_to_project_selection()
