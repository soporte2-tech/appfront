import streamlit as st
import google.generativeai as genai
import json
import re
import docx
from pypdf import PdfReader
import io
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.auth.transport.requests import Request
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
import time
import httplib2
import google_auth_httplib2
import os
import imgkit
# =============================================================================
#           CONFIGURACIÓN GLOBAL
# =============================================================================
st.set_page_config(layout="wide")

SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/userinfo.profile',
    'https://www.googleapis.com/auth/userinfo.email',
    'openid'
]
CLIENT_CONFIG = {
    "web": {
        "client_id": st.secrets["GOOGLE_CLIENT_ID"],
        "client_secret": st.secrets["GOOGLE_CLIENT_SECRET"],
        "redirect_uris": [st.secrets["GOOGLE_REDIRECT_URI"]],
        "auth_uri": "https://accounts.google.com/o/oauth2/auth",
        "token_uri": "https://oauth2.googleapis.com/token",
    }
}
ROOT_FOLDER_NAME = "ProyectosLicitaciones"

# =============================================================================
#           PROMPTS DE LA IA (CORREGIDOS Y CENTRALIZADOS)
# =============================================================================

PROMPT_PLIEGOS = """
Eres un consultor experto en licitaciones públicas, especializado en estructurar memorias técnicas para maximizar la puntuación. Tu conocimiento se basa ÚNICAMENTE en los archivos que te he proporcionado.
Tu misión es analizar los Pliegos (administrativos y técnicos) para diseñar un **índice jerárquico y estratégico** para la memoria técnica. Este índice debe responder perfectamente a todos los requisitos y, fundamentalmente, a los criterios de valoración.
## METODOLOGÍA DE ANÁLISIS OBLIGATORIA:
1.  **IDENTIFICAR APARTADOS PRINCIPALES:** Busca en los pliegos la sección de "CRITERIOS DE VALORACIÓN SUJETOS A JUICIO DE VALOR" (o similar). CADA UNO de estos criterios principales se convertirá en un **apartado de nivel superior** en tu estructura.
2.  **AGRUPAR SUBAPARTADOS LÓGICAMENTE:** Para cada apartado principal, busca en TODO el pliego los requisitos y detalles que correspondan a ese criterio. Estos detalles se convertirán en los **subapartados**.
3.  **GARANTIZAR COBERTURA TOTAL:** Asegúrate de que cada requisito relevante del pliego tenga su lugar en la estructura.
## REGLAS ESTRICTAS DE SALIDA:
0.  **LA JERARQUÍA ES CLAVE:** El objetivo es un índice bien estructurado. **Está prohibido generar una estructura con un único apartado principal.**
1.  **RESPUESTA EXCLUSIVAMENTE EN JSON:** Tu única salida debe ser un objeto JSON válido.
2.  **CLAVES PRINCIPALES FIJAS:** El JSON debe contener dos claves: "estructura_memoria" y "matices_desarrollo".
3.  **NUMERACIÓN JERÁRQUICA:** Usa numeración para cada apartado y subapartado (ej: "1. Título", "1.1. Subtítulo").
4.  **TÍTULOS FIELES AL PLIEGO:** Usa la terminología de los Pliegos.
5.  **CONTENIDO DE "matices_desarrollo":** Para CADA subapartado, las "indicaciones" deben incluir OBLIGATORIAMENTE:
    -   **Puntuación y Relevancia:** Menciona los puntos del criterio asociado.
    -   **Longitud Estimada:** Propón una longitud. NUNCA digas que no está especificado.
    -   **Contenido Detallado:** Explica qué información del pliego desarrollar.
    -   **Objetivo Estratégico:** Describe qué demostrar para conseguir la máxima puntuación.
    -   **Elementos Clave a Incluir:** Lista de puntos, tablas o gráficos que no pueden faltar.
## EJEMPLO DE ESTRUCTURA DE SALIDA OBLIGATORIA:
{
  "estructura_memoria": [
    {
      "apartado": "1. Solución Técnica y Metodología",
      "subapartados": ["1.1. Metodología de Trabajo", "1.2. Plan de Trabajo"]
    },
    {
      "apartado": "2. Calidad del Servicio y Mejoras",
      "subapartados": ["2.1. Actuaciones adicionales"]
    }
  ],
  "matices_desarrollo": [
    {
      "apartado": "1. Solución Técnica y Metodología",
      "subapartado": "1.1. Metodología de Trabajo",
      "indicaciones": "Este subapartado es clave para el criterio 'Calidad de la Propuesta Técnica', valorado con 40 puntos. Se recomienda una extensión de 8 páginas. Detallar la metodología agile-scrum, describiendo las fases del proyecto. Es fundamental incluir un diagrama de flujo. El objetivo es demostrar que nuestra metodología es robusta y flexible."
    }
  ]
}
"""

PROMPT_PREGUNTAS_TECNICAS_INDIVIDUAL = """
Actúa como un planificador de licitación. Tu objetivo será crear una propuesta de contenido ganadora para que tus compañeros la validen.
!! Tu respuesta debe centrarse EXCLUSIVAMENTE en el apartado proporcionado. No incluyas un índice general. Empieza directamente con el desarrollo. !!
Para el apartado proporcionado, responde a dos preguntas: "qué se debe incluir en este apartado" y "el contenido propuesto para ese apartado".
La primera pregunta ("Qué se debe incluir...") debe ser un resumen de lo que pide el pliego para ese apartado.
La segunda pregunta ("Contenido propuesto...") debe ser tu propuesta de contenido para obtener la mayor puntuación. Detállala ampliamente de manera esquemática. Desgrana el contenido general en preguntas más pequeñas y da respuestas detalladas.
REGLAS OBLIGATORIAS:
1.  **TONO PROFESIONAL E IMPERSONAL:** Redacta en tercera persona (ej: "El servicio se articula...", "La metodología implementada será...").
2.  **CONCRECIÓN ABSOLUTA:** Respalda cada afirmación con una acción concreta, una herramienta específica o una métrica medible.
3.  **ENFOQUE EN EL USUARIO FINAL:** Orienta el contenido a resolver los problemas del buyer persona.
4.  **LONGITUD CONTROLADA:** El desarrollo debe tener entre 6.000 y 8.000 palabras en total, distribuidas de forma equilibrada.
5.  **VALOR ESTRATÉGICO:** Enfócate en los resultados y el valor añadido.
6.  **ALINEACIÓN CON PLIEGO:** Justifica cada acción con su alineación con los requisitos del Pliego.
"""

PROMPT_REGENERACION = """
Actúas como un editor experto que refina una estructura JSON para una memoria técnica.
Te proporcionaré:
1.  Los documentos originales (Pliegos y/o plantilla).
2.  La estructura JSON que se generó en un primer intento.
3.  Las INSTRUCCIONES DE UN USUARIO con los cambios que desea.
Tu única tarea es generar una **NUEVA VERSIÓN MEJORADA** del objeto JSON que incorpore los cambios solicitados.
REGLAS OBLIGATORIAS:
-   **MANTÉN EL FORMATO ORIGINAL:** El JSON de salida debe ser válido y seguir la estructura original.
-   **INCORPORA EL FEEDBACK:** Aplica las correcciones del usuario a la nueva estructura.
-   **NO PIERDAS INFORMACIÓN:** Los apartados no mencionados por el usuario deben permanecer intactos.
-   **SÉ PRECISO:** No inventes nuevos apartados a menos que se te pida.
Genera únicamente el objeto JSON corregido.
"""

PROMPT_DESARROLLO = """
    **ATENCIÓN: REGLA CRÍTICA Y NO NEGOCIABLE**
    Tu única salida es el contenido final solicitado (texto en Markdown o un único bloque de código HTML). ESTÁ ABSOLUTAMENTE PROHIBIDO generar cualquier texto que analice, comente o critique tu propia salida. Frases como "Este código HTML...", "Puntos fuertes:", "Sugerencias:", "Con estas mejoras..." resultarán en un fallo inmediato. Debes actuar como el redactor final, no como un revisor de código.
    Actúa como un consultor experto redactando una memoria técnica. Tu tarea es crear el prompt para una IA que redactará un subapartado.
    El prompt debe ser exhaustivo y detallado.
    REGLAS CLAVE:
    1.  Crea UN SOLO prompt para el subapartado completo. Define un rango de palabras coherente.
    2.  Usa el CONTEXTO ADICIONAL (guion y docs de apoyo) como base principal.
    3.  Pide coherencia y referencias a lo dicho anteriormente para que la redacción parezca humana.
    4.  **REGLA DE ORO PARA ELEMENTOS VISUALES:** Si necesitas crear un elemento visual, DEBES generar un archivo HTML completo y auto-contenido, empezando con `<!DOCTYPE html>`.
    5.  **PROHIBIDO ANALIZAR O COMENTAR:** No escribas texto que analice tu propia salida.
    6.  **TONO Y ESTILO:** Mantén un tono profesional, impersonal (tercera persona) y concreto.
HERRAMIENTAS VISUALES A TU DISPOSICIÓN:
OPCIÓN A: PLANTILLA DE LISTA SIMPLE
```html
<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>Visual Element</title><style>@import url('https://fonts.googleapis.com/css2?family=Urbanist:wght@400;600;700&display=swap');body {{{{ font-family: 'Urbanist', sans-serif; background-color: #f0f2f5; display: flex; justify-content: center; align-items: center; padding: 20px; width: 800px; box-sizing: border-box; }}}} .card {{{{ background-color: white; border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); padding: 25px; width: 100%; max-width: 750px; border-top: 5px solid #0046C6; }}}} h2 {{{{ color: #0046C6; text-align: center; margin-top: 0; font-size: 24px; font-weight: 700; }}}} ul {{{{ list-style-type: none; padding: 0; }}}} li {{{{ display: flex; align-items: center; margin-bottom: 15px; font-size: 16px; color: #333; }}}} li::before {{{{ content: '✔'; color: #32CFAA; font-size: 20px; font-weight: bold; margin-right: 15px; }}}}</style></head><body><div class="card"><h2><!-- TÍTULO AQUÍ --></h2><ul><!-- LISTA DE ELEMENTOS AQUÍ --></ul></div></body></html>
OPCIÓN B: PLANTILLA DE INFOGRAFÍA MULTI-COLUMNA
<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>Infographic Element</title><style>@import url('https://fonts.googleapis.com/css2?family=Urbanist:wght@400;600;700&display=swap');body {{{{ font-family: 'Urbanist', sans-serif; display: flex; justify-content: center; align-items: center; padding: 20px; background-color: #f8f9fa; width: 800px; box-sizing: border-box; }}}} .container {{{{ background-color: #ffffff; border: 2px dashed #e0e0e0; border-radius: 15px; padding: 25px; width: 100%; box-shadow: 0 4px 12px rgba(0,0,0,0.05); }}}} h2 {{{{ color: #0046C6; text-align: center; font-size: 26px; font-weight: 700; margin-bottom: 25px; letter-spacing: -0.5px; }}}} .columns {{{{ display: flex; justify-content: space-around; gap: 20px; }}}} .column {{{{ flex: 1; text-align: center; padding: 15px; border-top: 4px solid; border-radius: 8px; background-color: #fdfdfd; }}}} .column-icon {{{{ width: 30px; height: 30px; border-radius: 50%; margin: 0 auto 15px auto; }}}} .column h3 {{{{ font-size: 16px; font-weight: 600; color: #333; margin-bottom: 10px; }}}} .column ul {{{{ list-style: none; padding: 0; margin: 0; text-align: left; }}}} .column li {{{{ font-size: 13px; color: #555; margin-bottom: 8px; line-height: 1.5; }}}} .color-1 {{{{ border-color: #FBC02D; }}}} .icon-1 {{{{ background-color: #FBC02D; }}}} .color-2 {{{{ border-color: #4CAF50; }}}} .icon-2 {{{{ background-color: #4CAF50; }}}} .color-3 {{{{ border-color: #90CAF9; }}}} .icon-3 {{{{ background-color: #90CAF9; }}}} .color-4 {{{{ border-color: #F44336; }}}} .icon-4 {{{{ background-color: #F44336; }}}}</style></head><body><div class="container"><h2><!-- TÍTULO --></h2><div class="columns"><div class="column color-1"><div class="column-icon icon-1"></div><h3><!-- Título Col 1 --></h3><ul><li><!-- Punto 1 --></li></ul></div><div class="column color-2"><div class="column-icon icon-2"></div><h3><!-- Título Col 2 --></h3><ul><li><!-- Punto 1 --></li></ul></div><div class="column color-3"><div class="column-icon icon-3"></div><h3><!-- Título Col 3 --></h3><ul><li><!-- Punto 1 --></li></ul></div></div></div></body></html>
Este es el subapartado para el que debes redactar los prompts:
- **Apartado Principal:** "{apartado_titulo}"
- **Subapartado a Redactar:** "{subapartado_titulo}"
Las instrucciones exactas de la plantilla para este subapartado son:
- **Indicaciones:** "{indicaciones}"
**REGLAS DE SALIDA:**
Tu respuesta DEBE ser SÓLO un único objeto JSON válido, con una única clave `"plan_de_prompts"` cuyo valor sea una lista de objetos. Cada objeto debe seguir esta estructura:
{{{{
  "apartado_referencia": "{apartado_titulo}",
  "subapartado_referencia": "{subapartado_titulo}",
  "prompt_id": "Un identificador único (ej: PROMPT_2_1_A)",
  "prompt_para_asistente": "La instrucción específica y detallada."
}}}}
=============================================================================
FUNCIONES DE AUTENTICACIÓN Y GOOGLE DRIVE
=============================================================================
def get_google_flow():
return Flow.from_client_config(
client_config=CLIENT_CONFIG,
scopes=SCOPES,
redirect_uri=st.secrets["GOOGLE_REDIRECT_URI"]
)
def get_credentials():
if 'credentials' in st.session_state and st.session_state.credentials:
creds = st.session_state.credentials
if not all(scope in creds.scopes for scope in SCOPES):
del st.session_state.credentials
go_to_landing()
st.rerun()
if 'credentials' in st.session_state and st.session_state.credentials:
    creds = st.session_state.credentials
    if creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
            st.session_state.credentials = creds
        except Exception:
            del st.session_state.credentials
            go_to_landing()
            st.rerun()
    return creds

if 'code' in st.query_params:
    try:
        flow = get_google_flow()
        flow.fetch_token(code=st.query_params['code'])
        st.session_state.credentials = flow.credentials
        st.query_params.clear()
        st.rerun()
    except Exception as e:
        st.error(f"Error al obtener el token: {e}")
        if 'credentials' in st.session_state: del st.session_state.credentials
        st.button("Reintentar inicio de sesión")
        st.stop()
        
return None
def build_drive_service(credentials):
try:
return build('drive', 'v3', credentials=credentials)
except HttpError as error:
st.error(f"No se pudo crear el servicio de Drive: {error}")
return None
=============================================================================
FUNCIONES AUXILIARES DE BACKEND
=============================================================================
def limpiar_respuesta_json(texto_sucio):
if not isinstance(texto_sucio, str): return ""
match = re.search(r'(?:json)?\s*(\{.*\})\s*', texto_sucio, re.DOTALL)
if match: return match.group(1).strip()
match = re.search(r'{.*}', texto_sucio, re.DOTALL)
if match: return match.group(0).strip()
return ""
def agregar_markdown_a_word(documento, texto_markdown):
patron_encabezado = re.compile(r'^(#+)\s+(.)')
patron_lista_numerada = re.compile(r'^\s\d+.\s+')
patron_lista_viñeta = re.compile(r'^\s*[*-]\s+')
def procesar_linea_con_negritas(parrafo, texto):
partes = re.split(r'(**.*?**)', texto)
for parte in partes:
if parte.startswith('') and parte.endswith(''):
parrafo.add_run(parte[2:-2]).bold = True
elif parte:
parrafo.add_run(parte)
for linea in texto_markdown.split('\n'):
    linea_limpia = linea.strip()
    if not linea_limpia: continue
    match_enc = patron_encabezado.match(linea_limpia)
    if match_enc:
        documento.add_heading(match_enc.group(2).strip(), level=min(len(match_enc.group(1)), 4))
        continue
    if patron_lista_numerada.match(linea_limpia):
        p = documento.add_paragraph(style='List Number')
        procesar_linea_con_negritas(p, patron_lista_numerada.sub('', linea_limpia))
    elif patron_lista_viñeta.match(linea_limpia):
        p = documento.add_paragraph(style='List Bullet')
        procesar_linea_con_negritas(p, patron_lista_viñeta.sub('', linea_limpia))
    else:
        p = documento.add_paragraph()
        procesar_linea_con_negritas(p, linea_limpia)
def mostrar_indice_desplegable(estructura_memoria):
if not estructura_memoria:
st.warning("No se encontró una estructura de memoria para mostrar.")
return
st.subheader("Índice Propuesto")
for seccion in estructura_memoria:
apartado_titulo = seccion.get("apartado", "Apartado sin título")
subapartados = seccion.get("subapartados", [])
with st.expander(f"{apartado_titulo}"):
if subapartados:
for sub in subapartados: st.markdown(f"- {sub}")
else: st.markdown("Este apartado no tiene subapartados definidos.")
def sanitize_json_string(json_str):
control_chars_regex = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')
return control_chars_regex.sub('', json_str)
def wrap_html_fragment(html_fragment):
if html_fragment.strip().startswith('<!DOCTYPE html>'):
return html_fragment
css_styles = """
@import url('https://fonts.googleapis.com/css2?family=Urbanist:wght@400;600;700&display=swap');
body { font-family: 'Urbanist', sans-serif; background-color: #f0f2f5; display: flex; justify-content: center; align-items: center; padding: 20px; width: 800px; box-sizing: border-box; }
.card { background-color: white; border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); padding: 25px; width: 100%; max-width: 750px; border-top: 5px solid #0046C6; }
h2 { color: #0046C6; text-align: center; margin-top: 0; font-size: 24px; font-weight: 700; }
ul { list-style-type: none; padding: 0; }
li { display: flex; align-items: center; margin-bottom: 15px; font-size: 16px; color: #333; }
li::before { content: '✔'; color: #32CFAA; font-size: 20px; font-weight: bold; margin-right: 15px; }
table { width: 100%; border-collapse: collapse; margin-top: 20px; font-size: 15px; }
th, td { padding: 12px 15px; border: 1px solid #ddd; text-align: left; }
th { background-color: #f5f5f5; font-weight: 600; color: #333; }
tr:nth-child(even) { background-color: #f9f9f9; }
"""
return f"""
<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>Visual Element</title><style>{css_styles}</style></head><body>{html_fragment}</body></html>
"""
def html_a_imagen(html_string, output_filename="temp_image.png"):
try:
path_wkhtmltoimage = os.popen('which wkhtmltoimage').read().strip()
if not path_wkhtmltoimage:
st.error("❌ El ejecutable 'wkhtmltoimage' no se encontró. Asegúrate de que 'wkhtmltopdf' está en tu packages.txt y que la app ha sido reiniciada.")
return None
config = imgkit.config(wkhtmltoimage=path_wkhtmltoimage)
options = {'format': 'png', 'encoding': "UTF-8", 'width': '800', 'quiet': ''}
imgkit.from_string(html_string, output_filename, config=config, options=options)
if os.path.exists(output_filename):
return output_filename
else:
st.warning(f"imgkit ejecutado pero el archivo '{output_filename}' no fue creado.")
return None
except Exception as e:
st.error(f"Ocurrió un error al convertir HTML a imagen: {e}")
st.code(f"Path de wkhtmltoimage intentado: {os.popen('which wkhtmltoimage').read().strip()}", language="bash")
return None
=============================================================================
NAVEGACIÓN Y GESTIÓN DE ESTADO
=============================================================================
if 'page' not in st.session_state: st.session_state.page = 'landing'
if 'credentials' not in st.session_state: st.session_state.credentials = None
if 'drive_service' not in st.session_state: st.session_state.drive_service = None
if 'selected_project' not in st.session_state: st.session_state.selected_project = None
def go_to_project_selection(): st.session_state.page = 'project_selection'
def go_to_landing(): st.session_state.page = 'landing'
def go_to_phase1(): st.session_state.page = 'phase_1'
def go_to_phase1_results(): st.session_state.page = 'phase_1_results'
def go_to_phase2(): st.session_state.page = 'phase_2'
def go_to_phase3(): st.session_state.page = 'phase_3'
def go_to_phase4(): st.session_state.page = 'phase_4'
def back_to_project_selection_and_cleanup():
keys_to_clear = ['generated_structure', 'uploaded_pliegos', 'selected_project', 'plan_status', 'generated_doc_buffer', 'generated_doc_filename']
for key in keys_to_clear:
if key in st.session_state: del st.session_state[key]
go_to_project_selection()
=============================================================================
PÁGINAS DE LA APLICACIÓN
=============================================================================
def landing_page():
"""Pantalla de bienvenida con inicio de sesión de Google."""
col1, col_center, col3 = st.columns([1, 2, 1])
with col_center:
st.image("https://raw.githubusercontent.com/soporte2-tech/appfront/main/imagen.png", width=150)
st.title("Asistente Inteligente para Memorias Técnicas")
st.header("Optimiza y acelera la creación de tus propuestas")
st.markdown("---")
st.info("Para empezar, necesitas dar permiso a la aplicación para que gestione los proyectos en tu Google Drive.")
flow = get_google_flow()
auth_url, _ = flow.authorization_url(prompt='consent')
st.link_button("🔗 Conectar con Google Drive", auth_url, use_container_width=True, type="primary")
def project_selection_page():
"""Página para seleccionar o crear un proyecto en Google Drive."""
st.header("Selección de Proyecto")
st.markdown("Elige un proyecto existente o crea uno nuevo para empezar.")
st.markdown("---")
if 'drive_service' not in st.session_state or not st.session_state.drive_service:
    st.session_state.drive_service = build_drive_service(st.session_state.credentials)

service = st.session_state.drive_service
if not service:
    st.error("No se pudo conectar con Google Drive. Por favor, intenta volver a la página de inicio y reconectar.")
    if st.button("← Volver a Inicio"):
        for key in ['credentials', 'drive_service']:
            if key in st.session_state: del st.session_state[key]
        go_to_landing()
        st.rerun()
    return

with st.spinner("Accediendo a tu Google Drive..."):
    root_folder_id = find_or_create_folder(service, ROOT_FOLDER_NAME)
    projects = list_project_folders(service, root_folder_id)

with st.container(border=True):
    st.subheader("1. Elige un proyecto existente")
    if not projects:
        st.info("Aún no tienes proyectos. Crea uno nuevo en el paso 2.")
    else:
        project_names = ["-- Selecciona un proyecto --"] + list(projects.keys())
        selected_name = st.selectbox("Selecciona tu proyecto:", project_names, label_visibility="collapsed")
        
        if st.button("Cargar Proyecto Seleccionado", type="primary", use_container_width=True):
            if selected_name != "-- Selecciona un proyecto --":
                st.session_state.selected_project = {"name": selected_name, "id": projects[selected_name]}
                st.toast(f"Proyecto '{selected_name}' cargado.")
                go_to_phase1()
                st.rerun()
            else:
                st.warning("Por favor, selecciona un proyecto de la lista.")

with st.container(border=True):
    st.subheader("2. O crea un nuevo proyecto")
    new_project_name = st.text_input("Nombre del nuevo proyecto", key="new_project_name_input", placeholder="Ej: Licitación Metro Madrid 2024", label_visibility="collapsed")
    if st.button("Crear y Empezar Nuevo Proyecto", use_container_width=True):
        if not new_project_name.strip():
            st.warning("Por favor, introduce un nombre para el proyecto.")
        elif new_project_name in projects:
            st.error("Ya existe un proyecto con ese nombre. Por favor, elige otro.")
        else:
            with st.spinner(f"Creando carpeta '{new_project_name}'..."):
                new_project_id = find_or_create_folder(service, new_project_name, parent_id=root_folder_id)
                st.session_state.selected_project = {"name": new_project_name, "id": new_project_id}
                st.success(f"¡Proyecto '{new_project_name}' creado!")
                go_to_phase1()
                st.rerun()
def phase_1_page(model):
"""Página de Fase 1: Carga de documentos y generación de índice."""
if not st.session_state.get('selected_project'):
st.warning("No se ha seleccionado ningún proyecto.")
go_to_project_selection()
st.rerun()
project_name = st.session_state.selected_project['name']
project_folder_id = st.session_state.selected_project['id']
service = st.session_state.drive_service

st.header("FASE 1: Análisis y Estructura")
st.info(f"Estás trabajando en el proyecto: **{project_name}**")

pliegos_folder_id = find_or_create_folder(service, "Pliegos", parent_id=project_folder_id)
document_files = get_files_in_project(service, pliegos_folder_id)

if document_files:
    st.success("Archivos encontrados en la carpeta 'Pliegos':")
    with st.container(border=True):
        for file in document_files:
            cols = st.columns([4, 1])
            cols[0].write(f"📄 **{file['name']}**")
            if cols[1].button("Eliminar", key=f"del_{file['id']}", type="secondary"):
                with st.spinner(f"Eliminando '{file['name']}'..."):
                    if delete_file_from_drive(service, file['id']):
                        st.toast(f"Archivo '{file['name']}' eliminado.")
                        st.rerun()
else:
    st.info("La carpeta 'Pliegos' está vacía. Sube los archivos base.")

with st.expander("Añadir o reemplazar documentación", expanded=not document_files):
    new_files_uploader = st.file_uploader("Arrastra aquí los Pliegos o Plantilla", type=['docx', 'pdf'], accept_multiple_files=True, key="new_files_uploader")
    if st.button("Guardar nuevos archivos en Drive"):
        if new_files_uploader:
            with st.spinner("Subiendo archivos..."):
                for file_obj in new_files_uploader:
                    upload_file_to_drive(service, file_obj, pliegos_folder_id)
                st.rerun()
        else:
            st.warning("Por favor, selecciona al menos un archivo.")

st.markdown("---")
st.subheader("Análisis y Generación de Índice")

docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)

col1, col2 = st.columns(2)
with col1:
    if st.button("Cargar último índice generado", use_container_width=True, disabled=not saved_index_id):
        with st.spinner("Cargando índice..."):
            index_content_bytes = download_file_from_drive(service, saved_index_id)
            st.session_state.generated_structure = json.loads(index_content_bytes.getvalue().decode('utf-8'))
            st.session_state.uploaded_pliegos = document_files
            go_to_phase1_results()
            st.rerun()
with col2:
    def handle_full_regeneration():
        with st.spinner("Analizando archivos y generando nuevo índice..."):
            try:
                downloaded_files = [{"mime_type": f['mimeType'], "data": download_file_from_drive(service, f['id']).getvalue()} for f in document_files]
                contenido_ia = [PROMPT_PLIEGOS] + downloaded_files
                response = model.generate_content(contenido_ia, generation_config=genai.GenerationConfig(response_mime_type="application/json"))
                json_limpio = limpiar_respuesta_json(response.text)
                if json_limpio:
                    st.session_state.generated_structure = json.loads(json_limpio)
                    st.session_state.uploaded_pliegos = document_files
                    go_to_phase1_results()
                else:
                    st.error("La IA devolvió una respuesta no válida.")
            except Exception as e:
                st.error(f"Error durante la regeneración: {e}")

    if st.button("Analizar y Generar Nuevo Índice", type="primary", use_container_width=True, disabled=not document_files):
        handle_full_regeneration()
        st.rerun()

st.markdown("---")
st.button("← Volver a Selección de Proyecto", on_click=back_to_project_selection_and_cleanup)
def phase_1_results_page(model):
"""Página para revisar, regenerar y aceptar el índice."""
st.header("FASE 1: Revisión de Resultados")
st.markdown("Revisa el índice. Puedes hacer ajustes, regenerarlo o aceptarlo para continuar.")
st.markdown("---")
st.button("← Volver a la gestión de archivos", on_click=go_to_phase1)
if 'generated_structure' not in st.session_state:
    st.warning("No se ha generado ninguna estructura."); return

def handle_regeneration_with_feedback():
    feedback_text = st.session_state.get("feedback_area", "")
    if not feedback_text:
        st.warning("Por favor, escribe tus indicaciones."); return
    with st.spinner("🧠 Incorporando tu feedback..."):
        try:
            service = st.session_state.drive_service
            contenido = [PROMPT_REGENERACION, f"--- INSTRUCCIONES DEL USUARIO ---\n{feedback_text}", f"--- ESTRUCTURA JSON ANTERIOR ---\n{json.dumps(st.session_state.generated_structure, indent=2)}"]
            if st.session_state.get('uploaded_pliegos'):
                for file_info in st.session_state.uploaded_pliegos:
                    contenido.append({"mime_type": file_info['mimeType'], "data": download_file_from_drive(service, file_info['id']).getvalue()})
            
            response = model.generate_content(contenido, generation_config=genai.GenerationConfig(response_mime_type="application/json"))
            json_limpio = limpiar_respuesta_json(response.text)
            if json_limpio:
                st.session_state.generated_structure = json.loads(json_limpio)
                st.toast("¡Estructura regenerada!")
                st.session_state.feedback_area = ""
            else:
                st.error("La IA no devolvió una estructura válida.")
        except Exception as e:
            st.error(f"Error en la regeneración: {e}")

with st.container(border=True):
    mostrar_indice_desplegable(st.session_state.generated_structure.get('estructura_memoria'))
    st.markdown("---")
    st.subheader("Validación y Siguiente Paso")
    st.text_area("Si necesitas cambios, indícalos aquí:", key="feedback_area", placeholder="Ej: 'Une los apartados 1.1 y 1.2 en uno solo.'")
    
    col1, col2 = st.columns(2)
    with col1:
        st.button("Regenerar con Feedback", on_click=handle_regeneration_with_feedback, use_container_width=True)
    with col2:
        # Aquí necesitaríamos re-ejecutar la lógica de la página anterior
        st.button("🔁 Regenerar Índice Entero", use_container_width=True, help="Vuelve a analizar los pliegos desde cero.")

    if st.button("Aceptar Índice y Pasar a Fase 2 →", type="primary", use_container_width=True):
        with st.spinner("Guardando índice y sincronizando carpetas..."):
            try:
                service = st.session_state.drive_service
                project_folder_id = st.session_state.selected_project['id']
                
                deleted_count = sync_guiones_folders_with_index(service, project_folder_id, st.session_state.generated_structure)
                if deleted_count > 0:
                    st.success(f"Limpieza completada: {deleted_count} carpetas obsoletas eliminadas.")

                docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
                json_bytes = json.dumps(st.session_state.generated_structure, indent=2).encode('utf-8')
                mock_file = io.BytesIO(json_bytes); mock_file.name = "ultimo_indice.json"; mock_file.type = "application/json"
                
                saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
                if saved_index_id: delete_file_from_drive(service, saved_index_id)
                upload_file_to_drive(service, mock_file, docs_app_folder_id)
                
                go_to_phase2()
                st.rerun()
            except Exception as e:
                st.error(f"Error en el guardado o sincronización: {e}")
def phase_2_page(model):
"""Centro de mando para la generación de guiones."""
st.header("FASE 2: Centro de Mando de Guiones")
st.markdown("Gestiona tus guiones de forma individual o selecciónalos para generarlos en lote.")
st.markdown("---")
service = st.session_state.drive_service
project_folder_id = st.session_state.selected_project['id']
if 'generated_structure' not in st.session_state:
    st.info("Sincronizando índice desde Google Drive...")
    try:
        docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)
        saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
        if saved_index_id:
            index_content_bytes = download_file_from_drive(service, saved_index_id).getvalue()
            st.session_state.generated_structure = json.loads(index_content_bytes.decode('utf-8'))
            st.rerun()
        else:
            st.warning("No se ha encontrado un índice guardado. Vuelve a la Fase 1.")
            if st.button("← Ir a Fase 1"): go_to_phase1(); st.rerun()
            return
    except Exception as e:
        st.error(f"Error al cargar el índice: {e}"); return

estructura = st.session_state.generated_structure.get('estructura_memoria', [])
matices_originales = st.session_state.generated_structure.get('matices_desarrollo', [])
matices_dict = {item.get('subapartado', ''): item for item in matices_originales if isinstance(item, dict)}

subapartados_a_mostrar = []
for seccion in estructura:
    for sub_titulo in seccion.get('subapartados', []):
        subapartados_a_mostrar.append(matices_dict.get(sub_titulo, {"subapartado": sub_titulo, "apartado": seccion.get('apartado')}))

with st.spinner("Sincronizando con Google Drive..."):
    guiones_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
    carpetas_existentes_resp = get_files_in_project(service, guiones_folder_id)
    nombres_carpetas_existentes = {f['name'] for f in carpetas_existentes_resp if f['mimeType'] == 'application/vnd.google-apps.folder'}

pending_keys = [matiz.get('subapartado') for matiz in subapartados_a_mostrar if re.sub(r'[\\/*?:"<>|]', "", matiz.get('subapartado')) not in nombres_carpetas_existentes]

def toggle_all_checkboxes():
    new_state = st.session_state.select_all_checkbox
    for key in pending_keys: st.session_state[f"cb_{key}"] = new_state

st.subheader("Generación de Borradores en Lote")
with st.container(border=True):
    col1, col2 = st.columns([1, 2])
    col1.checkbox("Seleccionar Todos / Ninguno", key="select_all_checkbox", on_change=toggle_all_checkboxes, disabled=not pending_keys)
    selected_keys = [key for key in pending_keys if st.session_state.get(f"cb_{key}")]
    
    if col2.button(f"🚀 Generar {len(selected_keys)} borradores seleccionados", type="primary", use_container_width=True, disabled=(not selected_keys)):
        progress_bar = st.progress(0, text="Iniciando generación...")
        items_to_generate = [matiz for matiz in subapartados_a_mostrar if matiz.get('subapartado') in selected_keys]
        for i, matiz in enumerate(items_to_generate):
            # Aquí iría la lógica de generación
            time.sleep(1) # Simulación
            progress_bar.progress((i + 1) / len(items_to_generate), text=f"Generando: {matiz.get('subapartado')}")
        progress_bar.progress(1.0, text="¡Completado!")
        st.success(f"{len(selected_keys)} borradores generados.")
        st.rerun()

st.markdown("---")
st.subheader("Gestión de Guiones de Subapartados")
for i, matiz in enumerate(subapartados_a_mostrar):
    sub_titulo = matiz.get('subapartado')
    if not sub_titulo: continue
    nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", sub_titulo)
    generado = nombre_limpio in nombres_carpetas_existentes

    with st.container(border=True):
        col1, col2 = st.columns([2, 1])
        with col1:
            if not generado:
                st.checkbox(f"**{sub_titulo}**", key=f"cb_{sub_titulo}")
            else:
                st.write(f"**{sub_titulo}**")
            st.caption(f"Estado: {'📄 Generado' if generado else '⚪ No Generado'}")
        with col2:
            if generado:
                st.button("Revisar en Drive", key=f"rev_{i}", use_container_width=True)
                st.button("🗑️ Borrar", key=f"del_{i}", use_container_width=True)
            else:
                st.button("Generar Borrador", key=f"gen_{i}", use_container_width=True)

st.markdown("---")
col_nav1, col_nav2 = st.columns(2)
col_nav1.button("← Volver a Revisión de Índice (F1)", on_click=go_to_phase1_results, use_container_width=True)
col_nav2.button("Ir a Plan de Prompts (F3) →", on_click=go_to_phase3, use_container_width=True)
def phase_3_page(model):
"""Página interactiva para generar, borrar, descargar y unificar planes de prompts."""
st.header("FASE 3: Centro de Mando de Prompts")
st.markdown("Gestiona los planes de prompts para cada subapartado.")
st.markdown("---")
service = st.session_state.drive_service
project_folder_id = st.session_state.selected_project['id']
docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)

def sync_plan_status_with_drive():
    with st.spinner("Sincronizando estado de los planes..."):
        guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
        carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
        st.session_state.plan_status = {}
        for nombre_carpeta, folder_id in carpetas_de_guiones.items():
            plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
            st.session_state.plan_status[nombre_carpeta] = {"guion_generado": True, "plan_id": plan_id}

if 'plan_status' not in st.session_state:
    sync_plan_status_with_drive()

if 'generated_structure' not in st.session_state:
    saved_index_id = find_file_by_name(service, "ultimo_indice.json", docs_app_folder_id)
    if saved_index_id:
        index_bytes = download_file_from_drive(service, saved_index_id).getvalue()
        st.session_state.generated_structure = json.loads(index_bytes.decode('utf-8'))
    else:
        st.warning("No se encontró el índice. Vuelve a la Fase 1.")
        if st.button("← Ir a Fase 1"): go_to_phase1(); st.rerun()
        return

estructura = st.session_state.generated_structure.get('estructura_memoria', [])
matices_map = {item.get('subapartado'): item for item in st.session_state.generated_structure.get('matices_desarrollo', [])}

subapartados_a_mostrar = []
for seccion in estructura:
    for sub_titulo in seccion.get('subapartados', []):
        subapartados_a_mostrar.append(matices_map.get(sub_titulo, {"subapartado": sub_titulo}))

def handle_individual_generation(matiz_info):
    sub_titulo = matiz_info.get("subapartado")
    nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", sub_titulo)
    with st.spinner(f"Generando plan para '{sub_titulo}'..."):
        try:
            # ... Lógica de generación de prompt y llamada a la IA ...
            # Al final:
            # new_file_id = upload_file_to_drive(...)
            st.session_state.plan_status[nombre_limpio]['plan_id'] = "temp_id" # Reemplazar con new_file_id
            st.toast(f"Plan para '{sub_titulo}' generado.")
        except Exception as e:
            st.error(f"Error generando prompts para '{sub_titulo}': {e}")

def handle_individual_deletion(titulo, plan_id):
    nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", titulo)
    with st.spinner(f"Eliminando plan para '{titulo}'..."):
        if delete_file_from_drive(service, plan_id):
            st.toast(f"Plan para '{titulo}' eliminado.")
            st.session_state.plan_status[nombre_limpio]['plan_id'] = None
        else:
            st.error("No se pudo eliminar el archivo.")
            
def handle_conjunto_generation():
    with st.spinner("Unificando todos los planes de prompts..."):
        try:
            guiones_main_folder_id = find_or_create_folder(service, "Guiones de Subapartados", parent_id=project_folder_id)
            carpetas_de_guiones = list_project_folders(service, guiones_main_folder_id)
            plan_conjunto_final = {"plan_de_prompts": []}
            
            for nombre_carpeta, folder_id in carpetas_de_guiones.items():
                plan_id = find_file_by_name(service, "prompts_individual.json", folder_id)
                if plan_id:
                    json_bytes = download_file_from_drive(service, plan_id).getvalue()
                    plan_individual_obj = json.loads(json_bytes.decode('utf-8'))
                    plan_conjunto_final["plan_de_prompts"].extend(plan_individual_obj.get("plan_de_prompts", []))

            if not plan_conjunto_final["plan_de_prompts"]:
                st.warning("No se encontraron planes para unificar."); return

            nombre_archivo = "plan_de_prompts_conjunto.json"
            json_bytes = json.dumps(plan_conjunto_final, indent=2, ensure_ascii=False).encode('utf-8')
            mock_file = io.BytesIO(json_bytes); mock_file.name = nombre_archivo; mock_file.type = "application/json"
            
            old_id = find_file_by_name(service, nombre_archivo, docs_app_folder_id)
            if old_id: delete_file_from_drive(service, old_id)
            upload_file_to_drive(service, mock_file, docs_app_folder_id)
            st.success(f"¡Plan conjunto generado con {len(plan_conjunto_final['plan_de_prompts'])} prompts!")
            st.balloons()
        except Exception as e:
            st.error(f"Error durante la unificación: {e}")

for matiz in subapartados_a_mostrar:
    sub_titulo = matiz.get("subapartado")
    nombre_limpio = re.sub(r'[\\/*?:"<>|]', "", sub_titulo)
    status = st.session_state.plan_status.get(nombre_limpio, {"guion_generado": False, "plan_id": None})
    
    with st.container(border=True):
        col1, col2 = st.columns([2, 1])
        with col1:
            st.write(f"**{sub_titulo}**")
            if not status["guion_generado"]:
                st.warning("⚠️ Guion no generado en Fase 2.")
            elif status["plan_id"]:
                st.success("✔️ Plan generado")
            else:
                st.info("⚪ Pendiente de generar plan")
        with col2:
            if status["guion_generado"]:
                if not status["plan_id"]:
                    st.button("Generar Plan", key=f"gen_{sub_titulo}", on_click=handle_individual_generation, args=(matiz,), use_container_width=True, type="primary")
                else:
                    st.button("Re-generar Plan", key=f"regen_{sub_titulo}", on_click=handle_individual_generation, args=(matiz,), use_container_width=True)
                    st.button("🗑️ Borrar Plan", key=f"del_{sub_titulo}", on_click=handle_individual_deletion, args=(sub_titulo, status['plan_id']), use_container_width=True)

st.markdown("---")
st.button("🚀 Unificar y Guardar Plan de Prompts Conjunto", on_click=handle_conjunto_generation, use_container_width=True, type="primary")
col_nav3_1, col_nav3_2 = st.columns(2)
col_nav3_1.button("← Volver a Guiones (F2)", on_click=go_to_phase2, use_container_width=True)
col_nav3_2.button("Ir a Redacción Final (F4) →", on_click=go_to_phase4, use_container_width=True)
def phase_4_page(model):
"""Página para ejecutar el plan de prompts y generar el documento Word final."""
st.header("FASE 4: Redacción y Ensamblaje Final")
st.markdown("Ejecuta el plan de prompts para generar el contenido de la memoria técnica.")
st.markdown("---")
service = st.session_state.drive_service
project_folder_id = st.session_state.selected_project['id']
docs_app_folder_id = find_or_create_folder(service, "Documentos aplicación", parent_id=project_folder_id)

plan_conjunto_id = find_file_by_name(service, "plan_de_prompts_conjunto.json", docs_app_folder_id)
if not plan_conjunto_id:
    st.warning("No se encontró el plan de prompts conjunto. Vuelve a la Fase 3.")
    if st.button("← Ir a Fase 3"): go_to_phase3(); st.rerun()
    return

try:
    json_bytes = download_file_from_drive(service, plan_conjunto_id).getvalue()
    plan_de_accion = json.loads(json_bytes.decode('utf-8'))
    lista_de_prompts = plan_de_accion.get("plan_de_prompts", [])
    st.success(f"✔️ Plan cargado con {len(lista_de_prompts)} prompts.")
except Exception as e:
    st.error(f"Error al cargar el plan de acción: {e}"); return

if 'generated_doc_buffer' not in st.session_state: st.session_state.generated_doc_buffer = None

button_text = "🔁 Volver a Generar Documento" if st.session_state.generated_doc_buffer else "🚀 Iniciar Redacción"
if st.button(button_text, type="primary", use_container_width=True):
    if not lista_de_prompts:
        st.warning("El plan de acción está vacío."); return

    with st.spinner("Iniciando redacción..."):
        progress_bar = st.progress(0, text="Configurando...")
        documento = docx.Document()
        chat_redaccion = model.start_chat()
        
        ultimo_apartado, ultimo_subapartado = "", ""
        
        for i, tarea in enumerate(lista_de_prompts):
            progress_bar.progress((i + 1) / len(lista_de_prompts), text=f"Procesando: {tarea.get('subapartado_referencia')}")
            
            apartado_actual = tarea.get("apartado_referencia", "")
            subapartado_actual = tarea.get("subapartado_referencia", "")
            if apartado_actual != ultimo_apartado:
                if ultimo_apartado: documento.add_page_break()
                documento.add_heading(apartado_actual, level=1)
                ultimo_apartado = apartado_actual; ultimo_subapartado = ""
            if subapartado_actual != ultimo_subapartado:
                documento.add_heading(subapartado_actual, level=2)
                ultimo_subapartado = subapartado_actual

            response = chat_redaccion.send_message(tarea.get("prompt_para_asistente"))
            agregar_markdown_a_word(documento, response.text)
            time.sleep(1)

        progress_bar.progress(1.0, text="Ensamblando documento...")
        project_name = st.session_state.selected_project['name']
        safe_name = re.sub(r'[\\/*?:"<>|]', "", project_name).replace(' ', '_')
        filename = f"Memoria_Tecnica_{safe_name}.docx"
        doc_io = io.BytesIO()
        documento.save(doc_io); doc_io.seek(0)
        
        st.session_state.generated_doc_buffer = doc_io
        st.session_state.generated_doc_filename = filename
        
        with st.spinner("Guardando en Google Drive..."):
            word_file = io.BytesIO(doc_io.getvalue()); word_file.name = filename; word_file.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            old_id = find_file_by_name(service, filename, docs_app_folder_id)
            if old_id: delete_file_from_drive(service, old_id)
            upload_file_to_drive(service, word_file, docs_app_folder_id)
        st.rerun()

if st.session_state.generated_doc_buffer:
    st.balloons()
    st.success("¡Tu documento está listo!")
    st.download_button(
        label="🎉 Descargar Memoria Técnica Final",
        data=st.session_state.generated_doc_buffer,
        file_name=st.session_state.generated_doc_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

st.markdown("---")
st.button("← Volver a Plan de Prompts (F3)", on_click=go_to_phase3, use_container_width=True)
=============================================================================
LÓGICA PRINCIPAL (ROUTER)
=============================================================================
credentials = get_credentials()
if not credentials:
landing_page()
else:
try:
genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel('gemini-1.5-pro-latest')
except Exception as e:
st.error(f"Error al configurar la API de Gemini. Verifica tu 'GEMINI_API_KEY'. Detalle: {e}")
st.stop()
if st.session_state.page in ['landing', 'project_selection']:
    project_selection_page()
elif st.session_state.page == 'phase_1':
    phase_1_page(model)
elif st.session_state.page == 'phase_1_results':
    phase_1_results_page(model)
elif st.session_state.page == 'phase_2':
    phase_2_page(model)
elif st.session_state.page == 'phase_3':
    phase_3_page(model)
elif st.session_state.page == 'phase_4':
    phase_4_page(model)
